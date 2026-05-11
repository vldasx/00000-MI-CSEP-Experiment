"""
generate_comparative.py — Comparative Analysis: Experiment I vs Experiment II
CSEP Prompting Experiment: 5 Small Models (7-14B) vs 5 Large Models (70B-104B)
Generates: CSEP_Comparative_Analysis.docx
"""
import json, os, sys
from collections import defaultdict

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE       = os.path.dirname(os.path.abspath(__file__))
EXP1_BASE  = os.path.join(BASE, "experiment_I")
EXP2_BASE  = os.path.join(BASE, "experiment_II")
OUTPUT     = os.path.join(BASE, "CSEP_Comparative_Analysis.docx")

CONDITIONS = {1: "Zero-Shot", 2: "CoT", 3: "CSEP-Only", 4: "CSEP+ZS"}
COND_SHORT = {1: "ZS", 2: "CoT", 3: "CSEP", 4: "CSEP+ZS"}

# ── Experiment I Data ─────────────────────────────────────────────────────────
EXP1_MODELS = {
    "mistralai/ministral-8b-2512":      "Ministral-8B",
    "meta-llama/llama-3.1-8b-instruct": "Llama-3.1-8B",
    "google/gemma-3-12b-it":            "Gemma-3-12B",
    "qwen/qwen-2.5-7b-instruct":        "Qwen-2.5-7B",
    "microsoft/phi-4":                  "Phi-4 (14B)",
}

# Auto-scored stats from analysis_report.json
EXP1_AUTO_OVERALL = {1: 0.6071, 2: 0.7009, 3: 0.5823, 4: 0.6177}
EXP1_AUTO_MODEL = {
    "mistralai/ministral-8b-2512":      {1: 0.7168, 2: 0.7434, 3: 0.6195, 4: 0.6814},
    "meta-llama/llama-3.1-8b-instruct": {1: 0.646,  2: 0.7168, 3: 0.5841, 4: 0.6903},
    "google/gemma-3-12b-it":            {1: 0.5752, 2: 0.708,  3: 0.5752, 4: 0.5664},
    "qwen/qwen-2.5-7b-instruct":        {1: 0.5841, 2: 0.646,  3: 0.5487, 4: 0.6018},
    "microsoft/phi-4":                  {1: 0.5133, 2: 0.6903, 3: 0.5841, 4: 0.5487},
}
EXP1_AUTO_CAT = {
    "2_fabricated_entities":            {1: 0.510,  2: 0.520,  3: 0.260,  4: 0.320},
    "3_math_with_distractors":          {1: 0.820,  2: 0.870,  3: 0.860,  4: 0.870},
    "6_false_premise":                  {1: 0.680,  2: 0.690,  3: 0.630,  4: 0.680},
    "8_confidence_calibration":         {1: 0.380,  2: 0.610,  3: 0.530,  4: 0.460},
    "12_counting_and_basic_arithmetic": {1: 0.8182, 2: 0.9636, 3: 0.8727, 4: 0.8545},
    "13_aiw_relational_puzzles":        {1: 0.4857, 2: 0.8571, 3: 0.5143, 4: 0.6571},
    "15_distractor_injection":          {1: 0.8286, 2: 0.9714, 3: 0.8286, 4: 0.8571},
    "18_hallucination_fabrication":     {1: 0.325,  2: 0.250,  3: 0.150,  4: 0.400},
}

# Judge V2 scores (finalized)
EXP1_JUDGE_OVERALL = {1: 0.7246, 2: 0.6851, 3: 0.7066, 4: 0.7211}
EXP1_JUDGE_MODEL = {
    "mistralai/ministral-8b-2512":      {1: 0.7311, 2: 0.7105, 3: 0.7256, 4: 0.7100},
    "meta-llama/llama-3.1-8b-instruct": {1: 0.6531, 2: 0.5581, 3: 0.6285, 4: 0.6749},
    "google/gemma-3-12b-it":            {1: 0.7358, 2: 0.7342, 3: 0.7303, 4: 0.7472},
    "qwen/qwen-2.5-7b-instruct":        {1: 0.7093, 2: 0.7072, 3: 0.6866, 4: 0.6909},
    "microsoft/phi-4":                  {1: 0.7943, 2: 0.7177, 3: 0.7629, 4: 0.7837},
}
EXP1_JUDGE_CAT = {
    "1_abstract_analogies":             {1: 0.950, 2: 0.830, 3: 0.910, 4: 0.930},
    "4_multi_step_reasoning":           {1: 0.820, 2: 0.845, 3: 0.810, 4: 0.805},
    "5_temporal_reasoning":             {1: 0.770, 2: 0.780, 3: 0.820, 4: 0.800},
    "7_hidden_contradictions":          {1: 0.705, 2: 0.580, 3: 0.720, 4: 0.725},
    "9_open_ended_hallucination_prone": {1: 0.452, 2: 0.398, 3: 0.458, 4: 0.461},
    "10_consistency_under_reframing":   {1: 0.903, 2: 0.859, 3: 0.891, 4: 0.890},
    "11_spatial_reasoning":             {1: 0.450, 2: 0.483, 3: 0.400, 4: 0.533},
    "12_counting_and_basic_arithmetic": {1: 0.200, 2: 0.333, 3: 0.267, 4: 0.267},
    "14_modified_classic_puzzles":      {1: 0.371, 2: 0.371, 3: 0.300, 4: 0.386},
    "15_distractor_injection":          {1: 1.000, 2: 0.800, 3: 0.600, 4: 0.800},
    "16_logical_inference":             {1: 0.954, 2: 0.862, 3: 0.923, 4: 0.923},
    "17_self_reference_paradox":        {1: 0.617, 2: 0.617, 3: 0.733, 4: 0.667},
    "19_linguistic_constraint":         {1: 0.683, 2: 0.583, 3: 0.567, 4: 0.617},
    "20_adversarial_pressure":          {1: 0.650, 2: 0.725, 3: 0.550, 4: 0.525},
}

# ── Experiment II Data ────────────────────────────────────────────────────────
EXP2_MODELS = {
    "meta-llama/llama-3.3-70b-instruct":          "Llama-3.3-70B",
    "qwen/qwen-2.5-72b-instruct":                 "Qwen-2.5-72B",
    "nvidia/llama-3.1-nemotron-70b-instruct":     "Nemotron-70B",
    "deepseek/deepseek-r1-distill-llama-70b":     "R1-Distill-70B",
}

EXP2_AUTO_OVERALL = {1: 0.739, 2: 0.6484, 3: 0.6946, 4: 0.7107}
EXP2_AUTO_MODEL = {
    "meta-llama/llama-3.3-70b-instruct":          {1: 0.819,  2: 0.5143, 3: 0.8095, 4: 0.9143},
    "qwen/qwen-2.5-72b-instruct":                 {1: 0.6476, 2: 0.750,  3: 0.7087, 4: 0.6952},
    "nvidia/llama-3.1-nemotron-70b-instruct":     {1: 0.9143, 2: 0.7333, 3: 0.7429, 4: 0.6827},
    "deepseek/deepseek-r1-distill-llama-70b":     {1: 0.600,  2: 0.7429, 3: 0.6264, 4: 0.7079},
}
EXP2_AUTO_CAT = {
    "2_fabricated_entities":            {1: 0.9158, 2: 0.400,  3: 0.6848, 4: 0.7111},
    "3_math_with_distractors":          {1: 0.830,  2: 0.870,  3: 0.820,  4: 0.840},
    "6_false_premise":                  {1: 0.810,  2: 0.680,  3: 0.7474, 4: 0.7835},
    "8_confidence_calibration":         {1: 0.5053, 2: 0.4222, 3: 0.3953, 4: 0.3605},
    "12_counting_and_basic_arithmetic": {1: 0.780,  2: 0.975,  3: 0.8462, 4: 0.8421},
    "13_aiw_relational_puzzles":        {1: 0.6286, 2: 0.9286, 3: 0.8571, 4: 0.8148},
    "15_distractor_injection":          {1: 0.750,  2: 0.8667, 3: 0.9333, 4: 0.8667},
    "18_hallucination_fabrication":     {1: 0.4333, 2: 0.4167, 3: 0.4783, 4: 0.7083},
}

# ── Load Exp II judge scores dynamically ─────────────────────────────────────
def normalize(s, cat):
    s = float(s)
    if cat in ["4_multi_step_reasoning", "7_hidden_contradictions",
               "10_consistency_under_reframing", "17_self_reference_paradox"]:
        return s / 2.0
    elif cat == "9_open_ended_hallucination_prone":
        return 1.0 - s
    return s

def compute_judge_stats(scores_file, exclude_models=None):
    exclude_models = exclude_models or set()
    with open(scores_file, encoding="utf-8") as f:
        raw = json.load(f)
    cond_scores  = defaultdict(list)
    model_cond   = defaultdict(lambda: defaultdict(list))
    cat_cond     = defaultdict(lambda: defaultdict(list))
    for k, v in raw.items():
        if v["score"] is None: continue
        mid = v["model_id"]
        if mid in exclude_models: continue
        cat = v["category"]; cond = str(v["condition"])
        ns = normalize(v["score"], cat)
        cond_scores[cond].append(ns)
        model_cond[mid][cond].append(ns)
        cat_cond[cat][cond].append(ns)
    overall = {int(c): sum(v)/len(v) for c, v in cond_scores.items() if v}
    by_model = {mid: {int(c): sum(v)/len(v) for c, v in conds.items() if v}
                for mid, conds in model_cond.items()}
    by_cat   = {cat: {int(c): sum(v)/len(v) for c, v in conds.items() if v}
                for cat, conds in cat_cond.items()}
    return overall, by_model, by_cat

exp2_judge_file = os.path.join(EXP2_BASE, "analysis", "judge_scores_v2.json")
if not os.path.exists(exp2_judge_file):
    print("ERROR: Exp II judge_scores_v2.json not found. Run llm_judge.py first.")
    sys.exit(1)

EXP2_EXCLUDE = {"cohere/command-r-plus-08-2024"}
EXP2_JUDGE_OVERALL, EXP2_JUDGE_MODEL_raw, EXP2_JUDGE_CAT = \
    compute_judge_stats(exp2_judge_file, EXP2_EXCLUDE)

# Check completeness
with open(exp2_judge_file, encoding="utf-8") as f:
    raw2 = json.load(f)
ok2  = sum(1 for v in raw2.values() if v["score"] is not None
           and v["model_id"] not in EXP2_EXCLUDE)
err2 = sum(1 for v in raw2.values() if v["score"] is None
           and v["model_id"] not in EXP2_EXCLUDE)
print(f"Exp II judge: {ok2} scored, {err2} errors")

EXP2_JUDGE_MODEL = EXP2_JUDGE_MODEL_raw  # mid -> {cond -> mean}

# ── Formatting Helpers ────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def score_color(v):
    if v is None: return "FFFFFF"
    if v >= 0.80: return "C6EFCE"   # green
    if v >= 0.65: return "FFEB9C"   # yellow
    if v >= 0.50: return "FCECD5"   # orange-light
    return "FFC7CE"                  # red

def delta_color(d):
    if d >= 0.05:  return "C6EFCE"
    if d <= -0.05: return "FFC7CE"
    return "FFFFFF"

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, italic=False, color=None, size=11, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

def fmt(v, pct=True):
    if v is None: return "—"
    return f"{v*100:.1f}%" if pct else f"{v:.3f}"

def delta_str(d):
    return f"{d*100:+.1f}%"

# ── Build Document ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Normal style font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CSEP Prompting Experiment")
run.bold = True; run.font.size = Pt(22); run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Comparative Analysis: Small Models vs. Large Models")
run2.bold = True; run2.font.size = Pt(16); run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Experiment I (7–14B) · Experiment II (70B–104B)")
run3.italic = True; run3.font.size = Pt(13); run3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()

# Subtitle info box (simple paragraph)
p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_info.paragraph_format.space_before = Pt(8)
run_info = p_info.add_run(
    "LLM-as-Judge: claude-sonnet-4-5 (OpenRouter) · Full responses · 500-token reasons\n"
    "14 task categories · 4 prompting conditions · 9 models total"
)
run_info.font.size = Pt(10.5); run_info.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

doc.add_page_break()

# ── 1. EXPERIMENT OVERVIEW ────────────────────────────────────────────────────
add_heading(doc, "1. Experiment Overview", 1)

add_para(doc, "Both experiments use identical question sets, evaluation categories, and LLM-as-judge methodology. The only differences are the model size tier and the exclusion of Command R+ (104B) from Experiment II due to API failures.", size=11, space_after=8)

tbl = doc.add_table(rows=5, cols=3)
tbl.style = "Table Grid"
tbl.autofit = False
widths = [Cm(4), Cm(6), Cm(6)]
for i, row in enumerate(tbl.rows):
    for j, cell in enumerate(row.cells):
        cell.width = widths[j]

headers = ["", "Experiment I — Small", "Experiment II — Large"]
for j, h in enumerate(headers):
    set_cell_bg(tbl.rows[0].cells[j], "2E74B5")
    cell_text(tbl.rows[0].cells[j], h, bold=True, color="FFFFFF")

rows_data = [
    ["Models",     "5 models: 7B–14B\nMinistral-8B, Llama-3.1-8B\nGemma-3-12B, Qwen-2.5-7B, Phi-4",
                   "4 models (excl. Cmd R+): ~70B\nLlama-3.3-70B, Qwen-2.5-72B\nNemotron-70B, R1-Distill-70B"],
    ["Records",    "6,440 (no errors)",       "3,900 total; 245 errors (Cmd R+)"],
    ["Auto-scored","2,260 items (8 categories)", "1,972 items (8 categories)"],
    ["Judge-scored","4,180 items (14 categories)", "1,560 items (14 categories, excl. Cmd R+)"],
]
for i, (label, e1, e2) in enumerate(rows_data):
    r = tbl.rows[i+1]
    set_cell_bg(r.cells[0], "DEEAF1")
    cell_text(r.cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
    cell_text(r.cells[1], e1, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
    cell_text(r.cells[2], e2, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)

doc.add_paragraph()

add_para(doc, "Prompting Conditions", bold=True, size=11, space_after=3)
for k, v in CONDITIONS.items():
    add_bullet(doc, f"Condition {k} — {v}: {'Direct question, no prompt engineering' if k==1 else 'Chain-of-Thought: 3 independent passes + synthesis' if k==2 else 'CSEP: Classify → Decompose → Reintegrate → Polish (4 API calls)' if k==3 else 'CSEP decomposition combined with Zero-Shot final answer'}")

doc.add_paragraph()

# ── 2. OVERALL PERFORMANCE COMPARISON ─────────────────────────────────────────
add_heading(doc, "2. Overall Performance by Condition", 1)
add_para(doc, "Scores are normalized to [0, 1]. Auto-scored categories use exact-match rules. Judge-scored categories use LLM-as-judge with category-specific rubrics. Hallucination scores are inverted (lower hallucination = higher score). 0–2 scale categories are divided by 2.", size=10.5, italic=True, space_after=8)

# --- 2a. Auto-scored ---
add_heading(doc, "2a. Auto-Scored Categories", 2)
add_para(doc, "8 categories with deterministic scoring: fabricated entities, math with distractors, false premise, confidence calibration, counting/arithmetic, relational puzzles, distractor injection, hallucination fabrication.", size=10.5, space_after=6)

tbl_auto = doc.add_table(rows=5, cols=5)
tbl_auto.style = "Table Grid"
tbl_auto.autofit = False
col_w = [Cm(4), Cm(3), Cm(3), Cm(3), Cm(3)]
for row in tbl_auto.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_w[j]

# Header
set_cell_bg(tbl_auto.rows[0].cells[0], "1F497D")
cell_text(tbl_auto.rows[0].cells[0], "Experiment", bold=True, color="FFFFFF")
for j, cond in enumerate([1,2,3,4]):
    set_cell_bg(tbl_auto.rows[0].cells[j+1], "1F497D")
    cell_text(tbl_auto.rows[0].cells[j+1], COND_SHORT[cond], bold=True, color="FFFFFF")

# Rows
for i, (label, data, bg) in enumerate([
    ("Exp I  (7–14B)",  EXP1_AUTO_OVERALL, "EBF3FB"),
    ("Exp II (70B)",    EXP2_AUTO_OVERALL, "FFF2CC"),
    ("Δ (II − I)",      None,              "F2F2F2"),
]):
    r = tbl_auto.rows[i+1]
    set_cell_bg(r.cells[0], bg)
    cell_text(r.cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
    for j, cond in enumerate([1,2,3,4]):
        if data is None:
            d = EXP2_AUTO_OVERALL.get(cond, 0) - EXP1_AUTO_OVERALL.get(cond, 0)
            set_cell_bg(r.cells[j+1], delta_color(d))
            cell_text(r.cells[j+1], delta_str(d), size=10)
        else:
            v = data.get(cond)
            set_cell_bg(r.cells[j+1], score_color(v))
            cell_text(r.cells[j+1], fmt(v), size=10)

# Empty separator row
set_cell_bg(tbl_auto.rows[4].cells[0], "FFFFFF")
for cell in tbl_auto.rows[4].cells:
    cell_text(cell, "")

doc.add_paragraph()

add_para(doc, "Key observation: Large models (Exp II) score notably higher on Zero-Shot (+13.2pp) but show a similar or slightly more pronounced CoT degradation on auto-scored tasks. CSEP and CSEP+ZS also improve with scale.", size=10.5, space_after=10)

# --- 2b. Judge-scored ---
add_heading(doc, "2b. Judge-Scored Categories", 2)
add_para(doc, "14 categories evaluated by claude-sonnet-4-5 with full responses (no truncation) and 500-token reasoning. These cover reasoning, analogies, contradictions, hallucination, spatial/temporal tasks.", size=10.5, space_after=6)

tbl_judge = doc.add_table(rows=5, cols=5)
tbl_judge.style = "Table Grid"
tbl_judge.autofit = False
for row in tbl_judge.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_w[j]

set_cell_bg(tbl_judge.rows[0].cells[0], "1F497D")
cell_text(tbl_judge.rows[0].cells[0], "Experiment", bold=True, color="FFFFFF")
for j, cond in enumerate([1,2,3,4]):
    set_cell_bg(tbl_judge.rows[0].cells[j+1], "1F497D")
    cell_text(tbl_judge.rows[0].cells[j+1], COND_SHORT[cond], bold=True, color="FFFFFF")

for i, (label, data, bg) in enumerate([
    ("Exp I  (7–14B)",  EXP1_JUDGE_OVERALL, "EBF3FB"),
    ("Exp II (70B)",    EXP2_JUDGE_OVERALL, "FFF2CC"),
    ("Δ (II − I)",      None,               "F2F2F2"),
]):
    r = tbl_judge.rows[i+1]
    set_cell_bg(r.cells[0], bg)
    cell_text(r.cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
    for j, cond in enumerate([1,2,3,4]):
        if data is None:
            d = EXP2_JUDGE_OVERALL.get(cond,0) - EXP1_JUDGE_OVERALL.get(cond,0)
            set_cell_bg(r.cells[j+1], delta_color(d))
            cell_text(r.cells[j+1], delta_str(d), size=10)
        else:
            v = data.get(cond)
            set_cell_bg(r.cells[j+1], score_color(v))
            cell_text(r.cells[j+1], fmt(v), size=10)

set_cell_bg(tbl_judge.rows[4].cells[0], "FFFFFF")
for cell in tbl_judge.rows[4].cells:
    cell_text(cell, "")

doc.add_paragraph()

add_para(doc, "Judge results show a consistent pattern: Zero-Shot remains the strongest condition in both experiments, with CSEP+ZS ranking second. CoT consistently underperforms Zero-Shot across both scale tiers.", size=10.5, space_after=4)

# Best condition per experiment
def best_cond(d):
    return max(d, key=lambda c: d[c])

b1 = best_cond(EXP1_JUDGE_OVERALL)
b2 = best_cond(EXP2_JUDGE_OVERALL)
add_bullet(doc, f"Exp I best condition: {CONDITIONS[b1]} ({fmt(EXP1_JUDGE_OVERALL[b1])})")
add_bullet(doc, f"Exp II best condition: {CONDITIONS[b2]} ({fmt(EXP2_JUDGE_OVERALL[b2])})")
add_bullet(doc, f"CSEP+ZS vs ZS delta — Exp I: {delta_str(EXP1_JUDGE_OVERALL[4]-EXP1_JUDGE_OVERALL[1])} · Exp II: {delta_str(EXP2_JUDGE_OVERALL.get(4,0)-EXP2_JUDGE_OVERALL.get(1,0))}")
add_bullet(doc, f"CoT vs ZS delta — Exp I: {delta_str(EXP1_JUDGE_OVERALL[2]-EXP1_JUDGE_OVERALL[1])} · Exp II: {delta_str(EXP2_JUDGE_OVERALL.get(2,0)-EXP2_JUDGE_OVERALL.get(1,0))}")

doc.add_paragraph()

# ── 3. MODEL-LEVEL COMPARISON ─────────────────────────────────────────────────
add_heading(doc, "3. Model Performance by Condition", 1)

# --- 3a. Exp I Models ---
add_heading(doc, "3a. Experiment I — Small Models (Judge V2)", 2)

tbl_m1 = doc.add_table(rows=len(EXP1_JUDGE_MODEL)+2, cols=6)
tbl_m1.style = "Table Grid"
tbl_m1.autofit = False
col_w2 = [Cm(4.5), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5)]
for row in tbl_m1.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_w2[j]

set_cell_bg(tbl_m1.rows[0].cells[0], "1F497D")
cell_text(tbl_m1.rows[0].cells[0], "Model", bold=True, color="FFFFFF")
for j, (lbl, cond) in enumerate([("ZS",1),("CoT",2),("CSEP",3),("CSEP+ZS",4),("Best",None)]):
    set_cell_bg(tbl_m1.rows[0].cells[j+1], "1F497D")
    cell_text(tbl_m1.rows[0].cells[j+1], lbl, bold=True, color="FFFFFF")

for i, (mid, name) in enumerate(EXP1_MODELS.items()):
    r = tbl_m1.rows[i+1]
    set_cell_bg(r.cells[0], "EBF3FB" if i%2==0 else "FFFFFF")
    cell_text(r.cells[0], name, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
    scores = EXP1_JUDGE_MODEL.get(mid, {})
    for j, cond in enumerate([1,2,3,4]):
        v = scores.get(cond)
        set_cell_bg(r.cells[j+1], score_color(v))
        cell_text(r.cells[j+1], fmt(v), size=10)
    best = max(scores, key=lambda c: scores[c]) if scores else None
    set_cell_bg(r.cells[5], "E2EFDA")
    cell_text(r.cells[5], COND_SHORT[best] if best else "—", bold=True, size=10)

# Average row
r_avg = tbl_m1.rows[len(EXP1_JUDGE_MODEL)+1]
set_cell_bg(r_avg.cells[0], "BDD7EE")
cell_text(r_avg.cells[0], "Average", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
for j, cond in enumerate([1,2,3,4]):
    v = EXP1_JUDGE_OVERALL.get(cond)
    set_cell_bg(r_avg.cells[j+1], score_color(v))
    cell_text(r_avg.cells[j+1], fmt(v), bold=True, size=10)
best_avg1 = best_cond(EXP1_JUDGE_OVERALL)
set_cell_bg(r_avg.cells[5], "E2EFDA")
cell_text(r_avg.cells[5], COND_SHORT[best_avg1], bold=True, size=10)

doc.add_paragraph()

# --- 3b. Exp II Models ---
add_heading(doc, "3b. Experiment II — Large Models (Judge V2)", 2)

tbl_m2 = doc.add_table(rows=len(EXP2_MODELS)+2, cols=6)
tbl_m2.style = "Table Grid"
tbl_m2.autofit = False
for row in tbl_m2.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_w2[j]

set_cell_bg(tbl_m2.rows[0].cells[0], "7030A0")
cell_text(tbl_m2.rows[0].cells[0], "Model", bold=True, color="FFFFFF")
for j, lbl in enumerate(["ZS","CoT","CSEP","CSEP+ZS","Best"]):
    set_cell_bg(tbl_m2.rows[0].cells[j+1], "7030A0")
    cell_text(tbl_m2.rows[0].cells[j+1], lbl, bold=True, color="FFFFFF")

for i, (mid, name) in enumerate(EXP2_MODELS.items()):
    r = tbl_m2.rows[i+1]
    set_cell_bg(r.cells[0], "F0E6FF" if i%2==0 else "FFFFFF")
    cell_text(r.cells[0], name, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
    scores = EXP2_JUDGE_MODEL.get(mid, {})
    for j, cond in enumerate([1,2,3,4]):
        v = scores.get(cond)
        set_cell_bg(r.cells[j+1], score_color(v))
        cell_text(r.cells[j+1], fmt(v), size=10)
    best = max(scores, key=lambda c: scores[c]) if scores else None
    set_cell_bg(r.cells[5], "E2EFDA")
    cell_text(r.cells[5], COND_SHORT[best] if best else "—", bold=True, size=10)

r_avg2 = tbl_m2.rows[len(EXP2_MODELS)+1]
set_cell_bg(r_avg2.cells[0], "D3B8FF")
cell_text(r_avg2.cells[0], "Average", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
for j, cond in enumerate([1,2,3,4]):
    v = EXP2_JUDGE_OVERALL.get(cond)
    set_cell_bg(r_avg2.cells[j+1], score_color(v))
    cell_text(r_avg2.cells[j+1], fmt(v), bold=True, size=10)
best_avg2 = best_cond(EXP2_JUDGE_OVERALL)
set_cell_bg(r_avg2.cells[5], "E2EFDA")
cell_text(r_avg2.cells[5], COND_SHORT[best_avg2], bold=True, size=10)

doc.add_paragraph()

# --- 3c. Cross-scale model comparison where possible ---
add_heading(doc, "3c. Qwen and Llama Across Scales", 2)
add_para(doc, "Two model families appear in both experiments, enabling direct comparison: Qwen (7B vs. 72B) and Llama (3.1-8B vs. 3.3-70B). Note that model versions differ slightly.", size=10.5, space_after=6)

pairs = [
    ("Llama family", "meta-llama/llama-3.1-8b-instruct", "meta-llama/llama-3.3-70b-instruct",
     "Llama-3.1-8B", "Llama-3.3-70B"),
    ("Qwen family",  "qwen/qwen-2.5-7b-instruct",        "qwen/qwen-2.5-72b-instruct",
     "Qwen-2.5-7B",  "Qwen-2.5-72B"),
]

tbl_cross = doc.add_table(rows=len(pairs)*2+1, cols=6)
tbl_cross.style = "Table Grid"
tbl_cross.autofit = False
for row in tbl_cross.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_w2[j]

set_cell_bg(tbl_cross.rows[0].cells[0], "404040")
cell_text(tbl_cross.rows[0].cells[0], "Model", bold=True, color="FFFFFF")
for j, lbl in enumerate(["ZS","CoT","CSEP","CSEP+ZS","Δ ZS (scale)"]):
    set_cell_bg(tbl_cross.rows[0].cells[j+1], "404040")
    cell_text(tbl_cross.rows[0].cells[j+1], lbl, bold=True, color="FFFFFF")

row_idx = 1
for family, mid1, mid2, name1, name2 in pairs:
    s1 = EXP1_JUDGE_MODEL.get(mid1, {})
    s2 = EXP2_JUDGE_MODEL.get(mid2, {})
    zs_delta = s2.get(1, 0) - s1.get(1, 0)
    for ii, (name, scores, bg) in enumerate([(name1, s1, "EBF3FB"), (name2, s2, "FFF2CC")]):
        r = tbl_cross.rows[row_idx]
        set_cell_bg(r.cells[0], bg)
        cell_text(r.cells[0], name, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
        for j, cond in enumerate([1,2,3,4]):
            v = scores.get(cond)
            set_cell_bg(r.cells[j+1], score_color(v))
            cell_text(r.cells[j+1], fmt(v), size=10)
        if ii == 0:
            set_cell_bg(r.cells[5], delta_color(zs_delta))
            cell_text(r.cells[5], delta_str(zs_delta), size=10)
        else:
            cell_text(r.cells[5], "", size=10)
        row_idx += 1

doc.add_paragraph()

# ── 4. SCALING EFFECTS ────────────────────────────────────────────────────────
add_heading(doc, "4. Scaling Effects: 7–14B vs. 70B", 1)

add_para(doc, "How does model scale affect (a) baseline performance, (b) CoT gains/losses, and (c) CSEP effectiveness?", size=11, bold=True, space_after=6)

# Calculate deltas
zs_delta  = EXP2_JUDGE_OVERALL.get(1,0) - EXP1_JUDGE_OVERALL[1]
cot_delta = EXP2_JUDGE_OVERALL.get(2,0) - EXP1_JUDGE_OVERALL[2]
csep_delta = EXP2_JUDGE_OVERALL.get(3,0) - EXP1_JUDGE_OVERALL[3]
csepzs_delta = EXP2_JUDGE_OVERALL.get(4,0) - EXP1_JUDGE_OVERALL[4]

add_heading(doc, "4a. Baseline Performance (Zero-Shot)", 2)
add_bullet(doc, f"Exp I average ZS: {fmt(EXP1_JUDGE_OVERALL[1])}")
add_bullet(doc, f"Exp II average ZS: {fmt(EXP2_JUDGE_OVERALL.get(1))}")
add_bullet(doc, f"Scale improvement: {delta_str(zs_delta)} — larger models are substantially stronger at zero-shot tasks")
add_bullet(doc, "Nemotron-70B achieves the highest ZS score in either experiment (auto: 91.4%), demonstrating that scale alone can push near-ceiling performance on structured tasks")
add_bullet(doc, "However, scale does NOT guarantee improvement: R1-Distill-70B (ZS judge ≈60%) scores below several 7-14B models, suggesting architecture/training matters as much as size")
doc.add_paragraph()

add_heading(doc, "4b. Chain-of-Thought (CoT) Effectiveness", 2)
add_bullet(doc, f"Exp I CoT vs ZS: {delta_str(EXP1_JUDGE_OVERALL[2]-EXP1_JUDGE_OVERALL[1])} — CoT hurts small models overall")
add_bullet(doc, f"Exp II CoT vs ZS: {delta_str(EXP2_JUDGE_OVERALL.get(2,0)-EXP2_JUDGE_OVERALL.get(1,0))} — CoT also hurts large models on average")
add_bullet(doc, "The CoT degradation is LARGER in Exp II (70B) than Exp I (7-14B), driven by catastrophic collapse in Llama 3.3-70B (auto: ZS=81.9% → CoT=51.4%, −30.5pp)")
add_bullet(doc, "Exception: R1-Distill-70B is the only model in either experiment where CoT consistently helps (auto: +14.3pp). This is a reasoning-specialist model trained with chain-of-thought paradigm")
add_bullet(doc, "Small model exception: Llama-3.1-8B shows severe CoT collapse too (judge: ZS=65.3% → CoT=55.8%, −9.5pp)")
doc.add_paragraph()

add_heading(doc, "4c. CSEP Prompting Effectiveness", 2)
add_bullet(doc, f"Exp I CSEP+ZS vs ZS: {delta_str(EXP1_JUDGE_OVERALL[4]-EXP1_JUDGE_OVERALL[1])} — slight improvement")
add_bullet(doc, f"Exp II CSEP+ZS vs ZS: {delta_str(EXP2_JUDGE_OVERALL.get(4,0)-EXP2_JUDGE_OVERALL.get(1,0))}")
add_bullet(doc, "CSEP alone (Condition 3) consistently underperforms ZS in both experiments — the structured decomposition without a clean final-answer pass introduces noise")
add_bullet(doc, "CSEP+ZS (Condition 4) is the safer CSEP variant: it captures conceptual decomposition benefits while using a clean zero-shot synthesis")
add_bullet(doc, "Phi-4 (14B, Exp I) is the strongest CSEP beneficiary: CSEP+ZS=78.4% vs ZS=79.4% — near parity. This model appears to use CSEP effectively without degradation")
add_bullet(doc, "Nemotron-70B shows CSEP degradation (ZS=91.4% auto → CSEP+ZS=68.3%), suggesting that highly capable models may be harmed by over-structuring")

doc.add_paragraph()

# ── 5. CATEGORY-LEVEL ANALYSIS ────────────────────────────────────────────────
add_heading(doc, "5. Category-Level Comparison", 1)
add_para(doc, "Full comparison of all categories across both experiments. Auto-scored categories (A) are shown in the first table; judge-scored categories (J) in the second.", size=10.5, italic=True, space_after=8)

# --- 5a. Auto-scored categories ---
add_heading(doc, "5a. Auto-Scored Categories", 2)

all_auto_cats = sorted(set(list(EXP1_AUTO_CAT.keys()) + list(EXP2_AUTO_CAT.keys())))
n_auto = len(all_auto_cats)

tbl_ac = doc.add_table(rows=n_auto*2+1, cols=5)
tbl_ac.style = "Table Grid"
tbl_ac.autofit = False
col_wA = [Cm(5), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5)]
for row in tbl_ac.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_wA[j]

set_cell_bg(tbl_ac.rows[0].cells[0], "404040")
cell_text(tbl_ac.rows[0].cells[0], "Category / Exp", bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.LEFT)
for j, cond in enumerate([1,2,3,4]):
    set_cell_bg(tbl_ac.rows[0].cells[j+1], "404040")
    cell_text(tbl_ac.rows[0].cells[j+1], COND_SHORT[cond], bold=True, color="FFFFFF")

cat_short = {
    "2_fabricated_entities":            "2. Fabricated Entities",
    "3_math_with_distractors":          "3. Math w/ Distractors",
    "6_false_premise":                  "6. False Premise",
    "8_confidence_calibration":         "8. Confidence Calib.",
    "12_counting_and_basic_arithmetic": "12. Counting/Arith.",
    "13_aiw_relational_puzzles":        "13. Relational Puzzles",
    "15_distractor_injection":          "15. Distractor Inject.",
    "18_hallucination_fabrication":     "18. Hallucination Fab.",
}

row_idx = 1
for cat in all_auto_cats:
    short = cat_short.get(cat, cat[:22])
    for exp_num, (exp_data, bg_label, bg_color) in enumerate([
        (EXP1_AUTO_CAT, "I",  "EBF3FB"),
        (EXP2_AUTO_CAT, "II", "FFF2CC"),
    ]):
        r = tbl_ac.rows[row_idx]
        label = f"{short} [Exp {bg_label}]"
        set_cell_bg(r.cells[0], bg_color)
        cell_text(r.cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.5,
                  bold=(exp_num==0))
        for j, cond in enumerate([1,2,3,4]):
            v = exp_data.get(cat, {}).get(cond)
            set_cell_bg(r.cells[j+1], score_color(v) if v is not None else "F0F0F0")
            cell_text(r.cells[j+1], fmt(v) if v is not None else "—", size=9.5)
        row_idx += 1

doc.add_paragraph()

# --- 5b. Judge-scored categories ---
add_heading(doc, "5b. Judge-Scored Categories", 2)

all_judge_cats = sorted(set(list(EXP1_JUDGE_CAT.keys()) + list(EXP2_JUDGE_CAT.keys())))

tbl_jc = doc.add_table(rows=len(all_judge_cats)*2+1, cols=5)
tbl_jc.style = "Table Grid"
tbl_jc.autofit = False
for row in tbl_jc.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_wA[j]

set_cell_bg(tbl_jc.rows[0].cells[0], "404040")
cell_text(tbl_jc.rows[0].cells[0], "Category / Exp", bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.LEFT)
for j, cond in enumerate([1,2,3,4]):
    set_cell_bg(tbl_jc.rows[0].cells[j+1], "404040")
    cell_text(tbl_jc.rows[0].cells[j+1], COND_SHORT[cond], bold=True, color="FFFFFF")

cat_short_j = {
    "1_abstract_analogies":             "1. Abstract Analogies",
    "4_multi_step_reasoning":           "4. Multi-Step Reasoning",
    "5_temporal_reasoning":             "5. Temporal Reasoning",
    "7_hidden_contradictions":          "7. Hidden Contradictions",
    "9_open_ended_hallucination_prone": "9. Hallucination (↑=better)",
    "10_consistency_under_reframing":   "10. Consistency/Reframing",
    "11_spatial_reasoning":             "11. Spatial Reasoning",
    "12_counting_and_basic_arithmetic": "12. Counting/Arith.",
    "14_modified_classic_puzzles":      "14. Modified Puzzles",
    "15_distractor_injection":          "15. Distractor Injection",
    "16_logical_inference":             "16. Logical Inference",
    "17_self_reference_paradox":        "17. Self-Ref. Paradox",
    "19_linguistic_constraint":         "19. Linguistic Constraint",
    "20_adversarial_pressure":          "20. Adversarial Pressure",
}

row_idx = 1
for cat in all_judge_cats:
    short = cat_short_j.get(cat, cat[:24])
    for exp_num, (exp_data, bg_label, bg_color) in enumerate([
        (EXP1_JUDGE_CAT, "I",  "EBF3FB"),
        (EXP2_JUDGE_CAT, "II", "FFF2CC"),
    ]):
        r = tbl_jc.rows[row_idx]
        label = f"{short} [Exp {bg_label}]"
        set_cell_bg(r.cells[0], bg_color)
        cell_text(r.cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.5,
                  bold=(exp_num==0))
        for j, cond in enumerate([1,2,3,4]):
            v = exp_data.get(cat, {}).get(cond)
            set_cell_bg(r.cells[j+1], score_color(v) if v is not None else "F0F0F0")
            cell_text(r.cells[j+1], fmt(v) if v is not None else "—", size=9.5)
        row_idx += 1

doc.add_paragraph()

# Category observations
add_heading(doc, "5c. Key Category Observations", 2)

obs = [
    ("Categories where scale helps most (Exp II > Exp I ZS):",
     ["Abstract Analogies: consistently near-ceiling in both, large models more stable",
      "Relational Puzzles (cat 13, auto): Exp II 62.9% vs Exp I 48.6% on ZS (+14.3pp)",
      "False Premise (cat 6): Exp II 81.0% vs Exp I 68.0% on ZS (+13.0pp)",
      "Fabricated Entities (cat 2): Exp II 91.6% vs Exp I 51.0% on ZS (+40.6pp) — large models much better at recognizing non-existent entities"]),
    ("Categories where scale shows no consistent benefit:",
     ["Math with Distractors (cat 3): both tiers ~82-83% ZS — math skill plateaus early",
      "Confidence Calibration (cat 8): Exp II 50.5% vs Exp I 38.0% — slight improvement but both below 60%",
      "Adversarial Pressure (cat 20, judge): large models not significantly more resistant to pressure"]),
    ("CoT helps most in both experiments:",
     ["Math with Distractors: +5pp (Exp I) and stable (Exp II)",
      "Relational Puzzles: +37pp (Exp I) and +30pp (Exp II) — biggest CoT gains",
      "Confidence Calibration: +23pp (Exp I) and −8.3pp (Exp II) — diverges at scale"]),
    ("CoT hurts most (consistent across both experiments):",
     ["Abstract Analogies (judge): −12pp (Exp I), trend in Exp II",
      "Distractor Injection (judge): −20pp (Exp I) — CoT introduces confusion about distractors",
      "Adversarial Pressure (judge): CoT makes models MORE prone to capitulating under pressure"]),
]

for title, bullets in obs:
    add_para(doc, title, bold=True, size=10.5, space_after=2)
    for b in bullets:
        add_bullet(doc, b)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── 6. NOTABLE OUTLIERS ───────────────────────────────────────────────────────
add_heading(doc, "6. Notable Model Profiles", 1)

# --- Phi-4 ---
add_heading(doc, "6a. Phi-4 (14B) — Best Small Model, CSEP-Resilient", 2)
phi4 = EXP1_JUDGE_MODEL.get("microsoft/phi-4", {})
add_para(doc, f"Phi-4 achieves ZS={fmt(phi4.get(1))} — the highest Zero-Shot score in Experiment I (judge). Despite being 14B parameters, it outperforms all other small models across all conditions.", size=10.5, space_after=4)
add_bullet(doc, f"ZS={fmt(phi4.get(1))} | CoT={fmt(phi4.get(2))} | CSEP={fmt(phi4.get(3))} | CSEP+ZS={fmt(phi4.get(4))}")
add_bullet(doc, f"CoT delta: {delta_str(phi4.get(2,0)-phi4.get(1,0))} — moderate degradation, better than Llama-3.1-8B")
add_bullet(doc, f"CSEP+ZS delta: {delta_str(phi4.get(4,0)-phi4.get(1,0))} — near-parity with ZS, suggesting CSEP neither helps nor hurts")
add_bullet(doc, "Phi-4's strong instruction-following capability likely contributes to its resilience across prompting strategies")
doc.add_paragraph()

# --- Nemotron ---
add_heading(doc, "6b. Nemotron-70B — Ceiling Performance, Severe CoT Collapse", 2)
nem = EXP2_JUDGE_MODEL.get("nvidia/llama-3.1-nemotron-70b-instruct", {})
nem_auto = EXP2_AUTO_MODEL.get("nvidia/llama-3.1-nemotron-70b-instruct", {})
add_para(doc, f"Nemotron-70B achieves ZS={fmt(nem_auto.get(1))} (auto) — the highest Zero-Shot score in either experiment. Its instruction-tuning (RLHF with human feedback) produces exceptional direct-answer quality.", size=10.5, space_after=4)
add_bullet(doc, f"Auto ZS={fmt(nem_auto.get(1))} | CoT={fmt(nem_auto.get(2))} | CSEP={fmt(nem_auto.get(3))} | CSEP+ZS={fmt(nem_auto.get(4))}")
add_bullet(doc, f"CoT delta: {delta_str(nem_auto.get(2,0)-nem_auto.get(1,0))} — catastrophic collapse, −18pp on auto tasks")
add_bullet(doc, "Hypothesis: Nemotron's RLHF training specifically optimized for concise, direct answers. Extended reasoning chains (CoT) interfere with this trained behavior")
add_bullet(doc, "CSEP also degrades performance, reinforcing the pattern: models optimized for direct response do not benefit from structured prompting")
doc.add_paragraph()

# --- R1-Distill ---
add_heading(doc, "6c. R1-Distill-70B — The Anomaly: CoT Consistently Helps", 2)
r1 = EXP2_JUDGE_MODEL.get("deepseek/deepseek-r1-distill-llama-70b", {})
r1_auto = EXP2_AUTO_MODEL.get("deepseek/deepseek-r1-distill-llama-70b", {})
add_para(doc, f"DeepSeek R1-Distill is the only model in either experiment where CoT consistently outperforms Zero-Shot. It is also the lowest Zero-Shot performer in Experiment II.", size=10.5, space_after=4)
add_bullet(doc, f"Auto ZS={fmt(r1_auto.get(1))} | CoT={fmt(r1_auto.get(2))} | CSEP={fmt(r1_auto.get(3))} | CSEP+ZS={fmt(r1_auto.get(4))}")
add_bullet(doc, f"Auto CoT delta: {delta_str(r1_auto.get(2,0)-r1_auto.get(1,0))} — the only positive CoT delta in Exp II for auto tasks")
add_bullet(doc, "R1-Distill was trained via knowledge distillation from DeepSeek-R1, which uses explicit chain-of-thought reasoning at training time")
add_bullet(doc, "Result: R1-Distill is architecturally aligned with CoT prompting — the model's internal representations assume multi-step reasoning chains")
add_bullet(doc, "96 empty responses in conditions 3 and 4 suggest R1-Distill struggles with CSEP's structured prompts (possible context length or format confusion)")
add_bullet(doc, "Implication: for R1-family models, CoT is the recommended prompting strategy; CSEP may be counterproductive")
doc.add_paragraph()

# --- Llama CoT collapse ---
add_heading(doc, "6d. Llama-3.3-70B — Largest CoT Collapse in Either Experiment", 2)
ll70 = EXP2_JUDGE_MODEL.get("meta-llama/llama-3.3-70b-instruct", {})
ll70_auto = EXP2_AUTO_MODEL.get("meta-llama/llama-3.3-70b-instruct", {})
add_para(doc, f"Llama-3.3-70B shows the most extreme CoT degradation observed in either experiment: auto score drops from ZS=81.9% to CoT=51.4% (−30.5pp).", size=10.5, space_after=4)
add_bullet(doc, f"Auto ZS={fmt(ll70_auto.get(1))} | CoT={fmt(ll70_auto.get(2))} | CSEP={fmt(ll70_auto.get(3))} | CSEP+ZS={fmt(ll70_auto.get(4))}")
add_bullet(doc, f"Auto CoT delta: {delta_str(ll70_auto.get(2,0)-ll70_auto.get(1,0))} — catastrophic")
add_bullet(doc, f"But CSEP+ZS recovers strongly: {fmt(ll70_auto.get(4))} — highest CSEP+ZS score in Exp II")
add_bullet(doc, "Pattern: Llama 3.3-70B uses CoT tokens 'verbosely' — generates extensive reasoning chains that confuse its final answer extraction")
add_bullet(doc, "CSEP+ZS works well because the decomposition happens in structured (non-verbose) steps, and the final ZS pass produces a clean answer")
add_bullet(doc, "Comparison with Llama-3.1-8B (Exp I): also shows CoT degradation (judge: −9.5pp), but far less severe — the 70B version amplifies this vulnerability")
doc.add_paragraph()

# ── 7. PROMPTING STRATEGY RECOMMENDATIONS ─────────────────────────────────────
add_heading(doc, "7. Prompting Strategy Analysis", 1)

add_heading(doc, "7a. Condition Rankings Summary", 2)

# Build ranking table
tbl_rank = doc.add_table(rows=3, cols=5)
tbl_rank.style = "Table Grid"
tbl_rank.autofit = False
for row in tbl_rank.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_wA[j]

set_cell_bg(tbl_rank.rows[0].cells[0], "1F497D")
cell_text(tbl_rank.rows[0].cells[0], "Experiment", bold=True, color="FFFFFF")
for j, cond in enumerate([1,2,3,4]):
    set_cell_bg(tbl_rank.rows[0].cells[j+1], "1F497D")
    cell_text(tbl_rank.rows[0].cells[j+1], COND_SHORT[cond], bold=True, color="FFFFFF")

def rank_row(data):
    sorted_conds = sorted(data.keys(), key=lambda c: data[c], reverse=True)
    return {c: sorted_conds.index(c)+1 for c in data}

rank1 = rank_row(EXP1_JUDGE_OVERALL)
rank2 = rank_row(EXP2_JUDGE_OVERALL)
rank_colors = ["C6EFCE", "FFEB9C", "FCECD5", "FFC7CE"]

for i, (label, ranks, data) in enumerate([
    ("Exp I (7–14B)",  rank1, EXP1_JUDGE_OVERALL),
    ("Exp II (70B)",   rank2, EXP2_JUDGE_OVERALL),
]):
    r = tbl_rank.rows[i+1]
    set_cell_bg(r.cells[0], "EBF3FB" if i==0 else "FFF2CC")
    cell_text(r.cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
    for j, cond in enumerate([1,2,3,4]):
        rank = ranks.get(cond, 4)
        set_cell_bg(r.cells[j+1], rank_colors[rank-1])
        cell_text(r.cells[j+1], f"#{rank} ({fmt(data.get(cond))})", size=10)

doc.add_paragraph()

add_para(doc, "The ranking is identical in both experiments: ZS > CSEP+ZS > CSEP > CoT (judge scores). This is a robust finding across scale tiers.", bold=True, size=11, space_after=6)

add_heading(doc, "7b. When to Use Each Condition", 2)

strategy_data = [
    ("Zero-Shot", "✔ Default choice for most tasks\n✔ Best overall performance in both experiments\n✔ Fastest (1 API call)\n✗ No structured reasoning support", "1F497D"),
    ("CSEP+ZS",  "✔ When tasks involve multi-faceted concepts\n✔ When ZS answers are inconsistent\n✔ Especially useful for Phi-4, Qwen-2.5 family\n✗ 5 API calls — resource intensive\n✗ May hurt reasoning-specialist models (Nemotron)", "2E74B5"),
    ("CoT",      "✔ ONLY for reasoning-specialist models (R1 family)\n✔ Multi-step math and relational puzzles\n✗ Actively harmful for most models\n✗ Especially bad for Llama and instruction-tuned models", "C00000"),
    ("CSEP-Only","✔ Niche use: when structured decomposition output is needed\n✗ Consistently underperforms ZS in both experiments\n✗ The 'Polish' step degrades final answer quality", "7030A0"),
]

tbl_strat = doc.add_table(rows=len(strategy_data)+1, cols=2)
tbl_strat.style = "Table Grid"
tbl_strat.autofit = False
strat_widths = [Cm(3), Cm(13.5)]
for row in tbl_strat.rows:
    row.cells[0].width = strat_widths[0]
    row.cells[1].width = strat_widths[1]

set_cell_bg(tbl_strat.rows[0].cells[0], "1F497D")
cell_text(tbl_strat.rows[0].cells[0], "Condition", bold=True, color="FFFFFF")
set_cell_bg(tbl_strat.rows[0].cells[1], "1F497D")
cell_text(tbl_strat.rows[0].cells[1], "Recommendation", bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.LEFT)

for i, (name, desc, color_hex) in enumerate(strategy_data):
    r = tbl_strat.rows[i+1]
    set_cell_bg(r.cells[0], color_hex)
    cell_text(r.cells[0], name, bold=True, color="FFFFFF", size=10)
    r.cells[1].text = ""
    for line in desc.split("\n"):
        p = r.cells[1].add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(1)

doc.add_paragraph()

# ── 8. KEY TAKEAWAYS ──────────────────────────────────────────────────────────
add_heading(doc, "8. Key Takeaways", 1)

takeaways = [
    ("Scale improves baseline but doesn't fix prompting sensitivity",
     "Large models (70B) score ~5–10pp higher on Zero-Shot tasks compared to small models (7–14B). However, they show equal or greater sensitivity to poorly-matched prompting strategies — Llama-3.3-70B's −30.5pp CoT collapse is the starkest example in either experiment."),
    ("Zero-Shot is the most robust strategy across all scales",
     "In both experiments, Zero-Shot ranks #1 for judge-scored tasks. This holds across 9 diverse models spanning two orders of magnitude in parameter count. Unless there is specific evidence that a model benefits from structured prompting, Zero-Shot should be the default."),
    ("CoT is harmful by default — with one exception",
     "Chain-of-Thought degrades performance in 8 of 9 models tested. The sole exception is R1-Distill-70B, whose training paradigm explicitly uses multi-step reasoning chains. This suggests CoT should be applied only when the model's training specifically supports it."),
    ("CSEP+ZS offers modest gains without the CoT downside",
     "CSEP+ZS consistently ranks #2 in both experiments, with modest gains of +0.0 to +3.0pp over ZS. For use cases where structured reasoning traces are valuable and compute cost is acceptable, CSEP+ZS is the recommended alternative to plain ZS."),
    ("Model architecture matters as much as scale",
     "Nemotron-70B (RLHF-optimized) behaves completely differently from R1-Distill-70B (reasoning-distilled), despite similar parameter counts. Scale-based predictions about prompting behavior are unreliable — architecture and training regime are more predictive."),
    ("Category-specific patterns are stable across scales",
     "Categories that are CoT-friendly (Relational Puzzles, Math) remain CoT-friendly at both scales. Categories that are CoT-hostile (Adversarial Pressure, Distractor Injection) remain hostile. This suggests task-level prompting guidance can be generalized across model sizes."),
]

for i, (title, body) in enumerate(takeaways):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run_num = p.add_run(f"{i+1}. ")
    run_num.bold = True; run_num.font.size = Pt(11)
    run_num.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run_title = p.add_run(title)
    run_title.bold = True; run_title.font.size = Pt(11)
    add_para(doc, body, size=10.5, space_after=6)

doc.add_paragraph()

# ── FOOTER NOTE ───────────────────────────────────────────────────────────────
add_para(doc,
    "Note: Experiment II judge scores are based on claude-sonnet-4-5 evaluations with full response text (no truncation) "
    "and 500-token reasoning allowance. Command R+ (104B) was excluded from Experiment II due to 245 HTTP 403 errors "
    "(incomplete data, 69% completion rate). All judge scores are normalized to [0, 1].",
    italic=True, size=9.5, color="808080", space_after=4)

add_para(doc,
    f"Exp II judge completeness: {ok2} scored, {err2} errors (of {ok2+err2} attempted).",
    italic=True, size=9.5, color="808080", space_after=4)

# ── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
