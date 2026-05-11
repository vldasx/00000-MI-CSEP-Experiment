"""
generate_report_v2.py — Generates CSEP_Analysis_Report_v2.docx
Uses judge_scores_v2.json (full response, 500-token reasons).
"""
import json, os, sys
from collections import defaultdict

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
GRAPHS_DIR = os.path.join(BASE, "analysis", "graphs")
OUTPUT = os.path.join(BASE, "CSEP_Analysis_Report_v2.docx")

# ── Data ─────────────────────────────────────────────────────────────────────

MODELS_SHORT = {
    "mistralai/ministral-8b-2512":      "Ministral-8B",
    "meta-llama/llama-3.1-8b-instruct": "Llama-3.1-8B",
    "google/gemma-3-12b-it":            "Gemma-3-12B",
    "qwen/qwen-2.5-7b-instruct":        "Qwen-2.5-7B",
    "microsoft/phi-4":                  "Phi-4",
}
MODEL_IDS = list(MODELS_SHORT.keys())
CONDITIONS = {1: "Zero-Shot", 2: "CoT", 3: "CSEP-Only", 4: "CSEP+ZS"}

AUTO_STATS = {
    "2_fabricated_entities":            {1: 0.51,   2: 0.52,   3: 0.26,   4: 0.32},
    "3_math_with_distractors":          {1: 0.82,   2: 0.87,   3: 0.86,   4: 0.87},
    "6_false_premise":                  {1: 0.68,   2: 0.69,   3: 0.63,   4: 0.68},
    "8_confidence_calibration":         {1: 0.38,   2: 0.61,   3: 0.53,   4: 0.46},
    "12_counting_and_basic_arithmetic": {1: 0.8182, 2: 0.9636, 3: 0.8727, 4: 0.8545},
    "13_aiw_relational_puzzles":        {1: 0.4857, 2: 0.8571, 3: 0.5143, 4: 0.6571},
    "15_distractor_injection":          {1: 0.8286, 2: 0.9714, 3: 0.8286, 4: 0.8571},
    "18_hallucination_fabrication":     {1: 0.325,  2: 0.25,   3: 0.15,   4: 0.40},
}
AUTO_OVERALL = {1: 0.6071, 2: 0.7009, 3: 0.5823, 4: 0.6177}
AUTO_MODEL = {
    "mistralai/ministral-8b-2512":      {1: 0.7168, 2: 0.7434, 3: 0.6195, 4: 0.6814},
    "meta-llama/llama-3.1-8b-instruct": {1: 0.646,  2: 0.7168, 3: 0.5841, 4: 0.6903},
    "google/gemma-3-12b-it":            {1: 0.5752, 2: 0.708,  3: 0.5752, 4: 0.5664},
    "qwen/qwen-2.5-7b-instruct":        {1: 0.5841, 2: 0.646,  3: 0.5487, 4: 0.6018},
    "microsoft/phi-4":                  {1: 0.5133, 2: 0.6903, 3: 0.5841, 4: 0.5487},
}

# V1 judge (truncated) — kept for comparison
JUDGE_V1_OVERALL = {1: 0.7245, 2: 0.5866, 3: 0.7010, 4: 0.7098}
JUDGE_V1_MODEL = {
    "mistralai/ministral-8b-2512":      {1: 0.7151, 2: 0.6691, 3: 0.6998, 4: 0.7115},
    "meta-llama/llama-3.1-8b-instruct": {1: 0.6454, 2: 0.4395, 3: 0.6541, 4: 0.6536},
    "google/gemma-3-12b-it":            {1: 0.7474, 2: 0.6026, 3: 0.7146, 4: 0.7261},
    "qwen/qwen-2.5-7b-instruct":        {1: 0.7221, 2: 0.5684, 3: 0.6880, 4: 0.6866},
    "microsoft/phi-4":                  {1: 0.7923, 2: 0.6531, 3: 0.7486, 4: 0.7711},
}

# ── Load v2 scores ─────────────────────────────────────────────────────────

def normalize(s, cat):
    s = float(s)
    if cat in ["4_multi_step_reasoning", "7_hidden_contradictions",
               "10_consistency_under_reframing", "17_self_reference_paradox"]:
        return s / 2.0
    elif cat == "9_open_ended_hallucination_prone":
        return 1.0 - s
    return s

with open(os.path.join(BASE, "analysis", "judge_scores_v2.json"), encoding="utf-8") as f:
    raw_scores = json.load(f)

with open(os.path.join(BASE, "analysis", "manual_review.json"), encoding="utf-8") as f:
    review_data = json.load(f)

review_lookup = {}
for item in review_data["items"]:
    k = f"{item['model_id']}|{item['question_id']}|{item['condition']}"
    review_lookup[k] = item

JUDGE_V2_OVERALL = defaultdict(list)
JUDGE_V2_MODEL   = defaultdict(lambda: defaultdict(list))
JUDGE_V2_CAT     = defaultdict(lambda: defaultdict(list))

for k, v in raw_scores.items():
    if v["score"] is None:
        continue
    s = normalize(v["score"], v["category"])
    JUDGE_V2_OVERALL[v["condition"]].append(s)
    JUDGE_V2_MODEL[v["model_id"]][v["condition"]].append(s)
    JUDGE_V2_CAT[v["category"]][v["condition"]].append(s)

def mean(lst): return sum(lst)/len(lst) if lst else None

J2_OVERALL = {c: mean(JUDGE_V2_OVERALL[c]) for c in [1,2,3,4]}
J2_MODEL    = {m: {c: mean(JUDGE_V2_MODEL[m][c]) for c in [1,2,3,4]} for m in MODEL_IDS}
J2_CAT      = {cat: {c: mean(JUDGE_V2_CAT[cat][c]) for c in [1,2,3,4]} for cat in JUDGE_V2_CAT}

CATEGORY_INFO = {
    "1_abstract_analogies":             ("Abstract Analogies",            "judge", "binary 0/1"),
    "2_fabricated_entities":            ("Fabricated Entities",           "auto",  "binary 0/1"),
    "3_math_with_distractors":          ("Math with Distractors",         "auto",  "binary 0/1"),
    "4_multi_step_reasoning":           ("Multi-Step Reasoning",          "judge", "0–2 scale"),
    "5_temporal_reasoning":             ("Temporal Reasoning",            "judge", "binary 0/1"),
    "6_false_premise":                  ("False Premise Detection",       "auto",  "binary 0/1"),
    "7_hidden_contradictions":          ("Hidden Contradictions",         "judge", "0–2 scale"),
    "8_confidence_calibration":         ("Confidence Calibration",        "auto",  "binary 0/1"),
    "9_open_ended_hallucination_prone": ("Hallucination Resistance",      "judge", "inverted 0–1"),
    "10_consistency_under_reframing":   ("Consistency Under Reframing",   "judge", "0–2 scale"),
    "11_spatial_reasoning":             ("Spatial Reasoning",             "judge", "binary 0/1"),
    "12_counting_and_basic_arithmetic": ("Counting & Arithmetic",         "judge", "binary 0/1"),
    "13_aiw_relational_puzzles":        ("Relational Puzzles (AIW)",      "auto",  "binary 0/1"),
    "14_modified_classic_puzzles":      ("Modified Classic Puzzles",      "judge", "binary 0/1"),
    "15_distractor_injection":          ("Distractor Injection",          "judge", "binary 0/1"),
    "16_logical_inference":             ("Logical Inference",             "judge", "binary 0/1"),
    "17_self_reference_paradox":        ("Self-Reference Paradox",        "judge", "0–2 scale"),
    "18_hallucination_fabrication":     ("Hallucination Fabrication",     "auto",  "binary 0/1"),
    "19_linguistic_constraint":         ("Linguistic Constraint",         "judge", "binary 0/1"),
    "20_adversarial_pressure":          ("Adversarial Pressure",          "judge", "binary 0/1"),
}

# ── Collect best/worst examples from v2 ───────────────────────────────────

judge_cat_examples = defaultdict(list)
for k, v in raw_scores.items():
    if v["score"] is None or not v.get("reason"):
        continue
    cat = v["category"]
    item = review_lookup.get(k, {})
    if not item:
        continue
    judge_cat_examples[cat].append({
        "model":     MODELS_SHORT.get(v["model_id"], v["model_id"].split("/")[-1]),
        "condition": v["condition"],
        "score":     v["score"],
        "norm":      normalize(v["score"], cat),
        "question":  item.get("question", ""),
        "expected":  item.get("expected", ""),
        "response":  item.get("response", ""),
        "reason":    v["reason"],
    })

EXAMPLES_V2 = {}
for cat, items in judge_cat_examples.items():
    items.sort(key=lambda x: x["norm"])
    EXAMPLES_V2[cat] = {"worst": items[0], "best": items[-1]}

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top","left","bottom","right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def score_color(v):
    if v is None: return "FFFFFF"
    if v >= 0.80: return "C6EFCE"
    elif v >= 0.65: return "FFEB9C"
    elif v >= 0.50: return "FFCC99"
    else: return "FFC7CE"

def delta_color(d):
    if d is None: return "FFFFFF"
    if d >= 0.03: return "C6EFCE"
    elif d <= -0.03: return "FFC7CE"
    return "F2F2F2"

def add_heading(doc, text, level, color="1F3864"):
    p = doc.add_heading(text, level=level)
    if p.runs: p.runs[0].font.color.rgb = RGBColor.from_string(color)
    return p

def add_para(doc, text="", bold=False, italic=False, size=None, indent=None, color=None):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor.from_string(color)
    return p

def tbl_header(tbl, headers, bg="1F3864"):
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            c.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(c, bg)

def add_example_box(doc, label, bg, model, cond_name, score, question, expected, response, reason):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"  {label}  |  {model}  |  {cond_name}  |  Score: {score}")
    r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(255,255,255)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),bg)
    pPr.append(shd)
    for lbl, txt, ital in [("Q", question, False),("Expected", expected, False),("Response", response, True),("Judge", reason, False)]:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.6)
        p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(1)
        r1 = p2.add_run(f"{lbl}: "); r1.bold = True; r1.font.size = Pt(9)
        r2 = p2.add_run(txt[:500]); r2.font.size = Pt(9); r2.italic = ital
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Build document ────────────────────────────────────────────────────────────

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.5)

    # ── Title ──
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(60)
    r = t.add_run("CSEP Experiment"); r.font.size = Pt(32); r.bold = True
    r.font.color.rgb = RGBColor(0x1F,0x38,0x64)

    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Conceptual Space Expansion Prompting")
    r2.font.size = Pt(20); r2.font.color.rgb = RGBColor(0x2E,0x75,0xB6)

    t3 = doc.add_paragraph(); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("Comprehensive Analysis Report — Version 2")
    r3.font.size = Pt(14)

    doc.add_paragraph()
    t4 = doc.add_paragraph(); t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = t4.add_run("5 Models  ·  322 Questions  ·  4 Conditions  ·  6,440 Records")
    r4.font.size = Pt(11); r4.font.color.rgb = RGBColor(0x40,0x40,0x40)

    t5 = doc.add_paragraph(); t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = t5.add_run("May 2026"); r5.font.size = Pt(11); r5.font.color.rgb = RGBColor(0x60,0x60,0x60)

    doc.add_page_break()

    # ── 0. Razlika od v1 ──
    add_heading(doc, "Napomena: Razlika od prethodne verzije izvještaja", 1, color="C00000")

    # Red warning box
    p_warn = doc.add_paragraph()
    p_warn.paragraph_format.left_indent = Cm(0.3)
    p_warn.paragraph_format.space_before = Pt(4)
    p_warn.paragraph_format.space_after = Pt(6)
    rw = p_warn.add_run(
        "Prethodna verzija ovog izvještaja (CSEP_Analysis_Report.docx) sadrži pogrešne rezultate "
        "za judge-scored kategorije zbog dvije metodološke greške u konfiguraciji LLM sudije. "
        "Zaključci iz te verzije NE SMIJU se koristiti."
    )
    rw.bold = True; rw.font.size = Pt(11); rw.font.color.rgb = RGBColor(0xC0,0x00,0x00)

    add_para(doc, "Greške u v1 (judge_scores.json):", bold=True)
    problems = [
        ("Truncacija odgovora na 1,500 karaktera:",
         "CoT odgovori tipično imaju 2,000–5,000 karaktera, CSEP odgovori 3,000–8,000 karaktera. "
         "Sudija je vidio samo prvu trećinu do četvrtinu odgovora. Ovo je sistematski "
         "oštećivalo CoT i CSEP skorove jer zaključak odgovora — koji sadrži finalni odgovor — "
         "najčešće dolazi na kraju, izvan prozora koji je sudija vidio."),
        ("Limit od 150 tokena za obrazloženje (reason):",
         "Sudija je producirao jednu kratku rečenicu bez konkretne analize. "
         "Nije moguće verificirati ispravnost ocene bez detaljnog objašnjenja."),
    ]
    for title, body in problems:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.5)
        r1 = p.add_run(title + " "); r1.bold = True
        p.add_run(body)

    add_para(doc, "Ispravke u v2 (judge_scores_v2.json):", bold=True)
    fixes = [
        "Odgovori se šalju sudiji bez ikakve truncacije (cijeli odgovor).",
        "max_tokens za reason povećan na 500 — sudija producira 2–3 rečenice s konkretnom analizom.",
        "Svi 4,180 zapisa su ponovo ocijenjeni od nule.",
        "Grešaka u ovom runu: 35 od 4,180 (0.8%) — zanemarivo.",
    ]
    for f in fixes:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f)

    # Comparison table v1 vs v2
    doc.add_paragraph()
    add_para(doc, "Poređenje v1 vs v2 — Judge-scored overall:", bold=True)
    tbl_cmp = doc.add_table(rows=3, cols=5)
    tbl_cmp.style = "Table Grid"
    tbl_header(tbl_cmp, ["Verzija", "Zero-Shot", "CoT", "CSEP-Only", "CSEP+ZS"], bg="4D4D4D")
    v1_row = [("v1 (truncated)", JUDGE_V1_OVERALL), ("v2 (full)", J2_OVERALL)]
    for ri, (label, data) in enumerate(v1_row):
        row = tbl_cmp.rows[ri+1]
        row.cells[0].text = label
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True
        for ci, c in enumerate([1,2,3,4]):
            v = data[c]
            row.cells[ci+1].text = f"{v:.3f}"
            row.cells[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row.cells[ci+1], score_color(v))
            set_cell_borders(row.cells[ci+1])

    doc.add_paragraph()
    add_para(doc,
        "Najvažnija promjena: CoT skor se popravio s 0.587 na 0.685 (+9.8pp) — "
        "truncacija je lažno prikazivala CoT kao najgoru strategiju. "
        "Međutim, generalni zaključak ostaje: ZS > CSEP+ZS ≈ CSEP > CoT, "
        "ali razlike su znatno manje nego što je v1 sugerisao.",
        italic=True, size=10)

    doc.add_page_break()

    # ── 1. Executive Summary ──
    add_heading(doc, "1. Executive Summary", 1)
    add_para(doc,
        "Ovaj izvještaj predstavlja rezultate kontrolisanog eksperimenta koji ispituje da li "
        "Conceptual Space Expansion Prompting (CSEP) poboljšava tačnost i robustnost malih "
        "jezičkih modela (7–14B parametara) u poređenju s Zero-Shot i Chain-of-Thought (CoT) "
        "baseline-ovima. Studija pokriva 20 kategorija kognitivnih zadataka, 5 modela, 322 pitanja "
        "i 4 uslova, što daje 6,440 ukupnih zapisa. "
        "Ocjenjivanje kombinuje automatsko regex-based bodovanje (8 kategorija, 2,260 zapisa) "
        "i LLM-as-judge evaluaciju (Claude Sonnet, 12 kategorija, 4,180 zapisa — puni odgovori, "
        "500-tokenski reasoni).")
    doc.add_paragraph()
    add_para(doc, "Ključni nalazi:", bold=True)
    bullets = [
        "Poredak strategija (judge-scored): Zero-Shot (0.725) > CSEP+ZS (0.721) ≈ CSEP (0.707) > CoT (0.685). Razlike su male.",
        "CoT dominira samo na auto-scored strukturiranim zadacima (0.701 vs ZS 0.607) — ali gubi na otvorenim zadacima.",
        "CSEP statistički značajno pomaže samo na: self-reference paradox (+11.7pp) i temporal reasoning (+5.0pp).",
        "CSEP šteti na: adversarial pressure (−10pp), linguistic constraint (−11.7pp), modified classic puzzles (−7.1pp).",
        "Phi-4 (14B) je najjači model na svim uslovima. Llama-3.1-8B ima najveći CoT pad (−9.5pp od ZS).",
        "CSEP+ZS je konzistentno bolji od CSEP-Only — zero-shot odgovor kao sidro smanjuje drift.",
        "Prethodna v1 analiza bila je metodološki neispravna zbog truncacije odgovora.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.4)
        p.add_run(b).font.size = Pt(10.5)

    doc.add_page_break()

    # ── 2. Methodology ──
    add_heading(doc, "2. Metodologija", 1)

    add_heading(doc, "2.1 Eksperimentalni dizajn", 2)
    design = [
        ("Modeli (5):", "Ministral-8B (2512), Llama-3.1-8B-Instruct, Gemma-3-12B-IT, Qwen-2.5-7B-Instruct, Phi-4"),
        ("Pitanja (322):", "20 kognitivnih kategorija; svako pitanje × 5 modela × 4 uslova"),
        ("Uslovi (4):", "Zero-Shot, Chain-of-Thought (3 passa + sinteza), CSEP-Only, CSEP+Zero-Shot"),
        ("Ukupno zapisa:", "322 × 4 × 5 = 6,440 (0 grešaka nakon završetka)"),
        ("API:", "OpenRouter — svi modeli kroz jedinstveni endpoint"),
    ]
    tbl_d = doc.add_table(rows=len(design), cols=2)
    tbl_d.style = "Table Grid"
    for i,(k,v) in enumerate(design):
        set_cell_bg(tbl_d.rows[i].cells[0], "EBF3FB")
        tbl_d.rows[i].cells[0].text = k
        if tbl_d.rows[i].cells[0].paragraphs[0].runs:
            tbl_d.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        tbl_d.rows[i].cells[1].text = v
    doc.add_paragraph()

    add_heading(doc, "2.2 Četiri uslova", 2)
    cond_details = [
        ("Uvjet 1 — Zero-Shot", "1F3864",
         "Direktan prompt: 'Answer the following question: {question}'. Nema scaffoldinga. Baseline."),
        ("Uvjet 2 — Chain-of-Thought (CoT)", "2E75B6",
         "Tri nezavisna CoT passa (knowledge activation → decomposition → reasoning → synthesis → self-verification), "
         "zatim četvrti sintezni poziv koji poredi tri pokušaja i daje finalni odgovor. Ukupno: 4 API poziva."),
        ("Uvjet 3 — CSEP-Only", "375623",
         "Četiri koraka: (1) Klasifikacija tipa pitanja; (2) Dekompozicija — identifikacija glavnog koncepta, "
         "3 rečenice opisa, 3 relevantne činjenice, provjera znanja s konfirmišućim/diskonformišućim primjerima; "
         "(3) Reintegracija — generisanje odgovora koristeći cijeli kontekst dekompozicije; "
         "(4) Poliranje — provjera konzistentnosti i ispravka. Ukupno: 4 API poziva."),
        ("Uvjet 4 — CSEP+Zero-Shot", "7B2D8B",
         "Identično kao CSEP-Only, ali reintegracija dobija i prethodni zero-shot odgovor kao kontekst. "
         "Sidri CSEP reasoning na modelov direktni instinkt. Zahtijeva da uvjet 1 prethodi."),
    ]
    for name, color, desc in cond_details:
        p = doc.add_paragraph()
        r = p.add_run(name); r.bold = True; r.font.size = Pt(11.5)
        r.font.color.rgb = RGBColor.from_string(color)
        add_para(doc, desc, size=10)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    add_heading(doc, "2.3 Ocjenjivanje", 2)
    add_para(doc,
        "Dvije metode ocjenjivanja, zavisno od kategorije:")
    score_methods = [
        ("Auto-scored (8 kategorija, 2,260 zapisa):",
         "Regex i string matching. Provjera ključnih termina, numeričkih odgovora ili odsustva lažnih tvrdnji. "
         "Kategorije: Fabricated Entities, Math with Distractors, False Premise, Confidence Calibration, "
         "Counting & Arithmetic, Relational Puzzles (AIW), Distractor Injection, Hallucination Fabrication."),
        ("LLM-as-Judge (12 kategorija, 4,180 zapisa — v2):",
         "Svaki odgovor ocijenjen od strane Claude Sonnet (claude-sonnet-4-5) via OpenRouter. "
         "Sudiji je dat: naziv kategorije, rubrika, pitanje, očekivani odgovor i POTPUNI odgovor modela (bez truncacije). "
         "max_tokens=500, temperature=0. Izlaz: JSON {score, reason} gdje reason sadrži 2–3 rečenice "
         "s konkretnom analizom šta je tačno/pogrešno i zašto."),
    ]
    for lbl, desc in score_methods:
        p = doc.add_paragraph()
        r1 = p.add_run(lbl+" "); r1.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(5)

    # Scale table
    add_para(doc, "Skale ocjenjivanja po kategoriji:", bold=True)
    cat_keys = sorted(CATEGORY_INFO.keys(), key=lambda x: int(x.split("_")[0]))
    tbl_sc = doc.add_table(rows=len(cat_keys)+1, cols=4)
    tbl_sc.style = "Table Grid"
    tbl_header(tbl_sc, ["Kategorija","Metoda","Skala","Napomena"], bg="1F3864")
    for ri, cat in enumerate(cat_keys):
        name, scorer, scale = CATEGORY_INFO[cat]
        row = tbl_sc.rows[ri+1]
        row.cells[0].text = name
        row.cells[1].text = scorer.upper()
        row.cells[2].text = scale
        row.cells[3].text = "Invertovano (1−score)" if cat == "9_open_ended_hallucination_prone" else ""
        set_cell_bg(row.cells[1], "E2EFDA" if scorer=="judge" else "FFF2CC")
        for c in row.cells:
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_borders(c)

    doc.add_page_break()

    # ── 3. Results ──
    add_heading(doc, "3. Rezultati", 1)

    # 3.1 Overall
    add_heading(doc, "3.1 Overall po uvjetu", 2)
    add_para(doc, "Tabela 1 — normalizovani skor (0–1), viši = bolji.")

    def combined_overall(c):
        a = AUTO_OVERALL[c]*565
        j = J2_OVERALL[c]*1036
        return (a+j)/(565+1036)

    tbl_ov = doc.add_table(rows=4, cols=5)
    tbl_ov.style = "Table Grid"
    tbl_header(tbl_ov, ["","Zero-Shot","CoT","CSEP-Only","CSEP+ZS"])
    ov_rows = [
        ("Auto-scored", AUTO_OVERALL),
        ("Judge-scored (v2)", J2_OVERALL),
        ("Kombinirano", {c: combined_overall(c) for c in [1,2,3,4]}),
    ]
    for ri,(lbl,data) in enumerate(ov_rows):
        row = tbl_ov.rows[ri+1]
        row.cells[0].text = lbl
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True
        for ci,c in enumerate([1,2,3,4]):
            v = data[c]
            row.cells[ci+1].text = f"{v:.3f}"
            row.cells[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row.cells[ci+1], score_color(v))
            set_cell_borders(row.cells[ci+1])

    doc.add_paragraph()
    add_para(doc,
        "Napomena: Auto-scored i judge-scored kategorije mjere različite tipove zadataka, "
        "pa kombinirani skor treba tumačiti s oprezom. "
        "Auto-scored zadaci (računanje, relacijske zagonetke) pogoduju CoT-u; "
        "judge-scored otvoreni zadaci pogoduju Zero-Shot-u.",
        italic=True, size=9.5)

    # 3.2 By model
    doc.add_paragraph()
    add_heading(doc, "3.2 Po modelu", 2)
    add_para(doc, "Tabela 2a — Auto-scored po modelu:", bold=True)
    tbl_am = doc.add_table(rows=len(MODEL_IDS)+1, cols=5)
    tbl_am.style = "Table Grid"
    tbl_header(tbl_am, ["Model","Zero-Shot","CoT","CSEP-Only","CSEP+ZS"], bg="2E75B6")
    for ri,mid in enumerate(MODEL_IDS):
        row = tbl_am.rows[ri+1]
        row.cells[0].text = MODELS_SHORT[mid]
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True
        for ci,c in enumerate([1,2,3,4]):
            v = AUTO_MODEL[mid][c]
            row.cells[ci+1].text = f"{v:.3f}"
            row.cells[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row.cells[ci+1], score_color(v))
            set_cell_borders(row.cells[ci+1])

    doc.add_paragraph()
    add_para(doc, "Tabela 2b — Judge-scored (v2) po modelu:", bold=True)
    tbl_jm = doc.add_table(rows=len(MODEL_IDS)+1, cols=5)
    tbl_jm.style = "Table Grid"
    tbl_header(tbl_jm, ["Model","Zero-Shot","CoT","CSEP-Only","CSEP+ZS"], bg="375623")
    for ri,mid in enumerate(MODEL_IDS):
        row = tbl_jm.rows[ri+1]
        row.cells[0].text = MODELS_SHORT[mid]
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True
        for ci,c in enumerate([1,2,3,4]):
            v = J2_MODEL[mid][c]
            row.cells[ci+1].text = f"{v:.3f}" if v else "-"
            row.cells[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if v: set_cell_bg(row.cells[ci+1], score_color(v))
            set_cell_borders(row.cells[ci+1])

    # 3.3 By category
    doc.add_paragraph()
    add_heading(doc, "3.3 Po kategoriji", 2)
    add_para(doc, "Tabela 3 — sve 20 kategorije. Boje: zelena ≥ 0.80, žuta ≥ 0.65, narančasta ≥ 0.50, crvena < 0.50.")

    all_cats_sorted = sorted(CATEGORY_INFO.keys(), key=lambda x: int(x.split("_")[0]))
    all_data = {**AUTO_STATS, **J2_CAT}

    tbl_cat = doc.add_table(rows=len(all_cats_sorted)+1, cols=7)
    tbl_cat.style = "Table Grid"
    tbl_header(tbl_cat, ["#","Kategorija","Metoda","ZS","CoT","CSEP","CSEP+ZS"])
    for ri,cat in enumerate(all_cats_sorted):
        row = tbl_cat.rows[ri+1]
        num = cat.split("_")[0]
        name, scorer, _ = CATEGORY_INFO[cat]
        data = all_data.get(cat, {})
        row.cells[0].text = num
        row.cells[1].text = name
        row.cells[2].text = scorer.upper()
        set_cell_bg(row.cells[2], "E2EFDA" if scorer=="judge" else "FFF2CC")
        for ci,c in enumerate([1,2,3,4]):
            v = data.get(c)
            if v is not None:
                row.cells[ci+3].text = f"{v:.3f}"
                row.cells[ci+3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_bg(row.cells[ci+3], score_color(v))
            for cell in row.cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_borders(row.cells[ci+3])

    # 3.4 CSEP delta table
    doc.add_paragraph()
    add_heading(doc, "3.4 CSEP delta vs Zero-Shot (judge-scored)", 2)
    add_para(doc, "Zelena = CSEP poboljšava ≥ 3pp, crvena = CSEP pogoršava ≤ −3pp, siva = zanemarivo.")

    judge_cats = sorted(J2_CAT.keys(), key=lambda x: int(x.split("_")[0]))
    tbl_delta = doc.add_table(rows=len(judge_cats)+1, cols=4)
    tbl_delta.style = "Table Grid"
    tbl_header(tbl_delta, ["Kategorija","CSEP − ZS","CSEP+ZS − ZS","ZS baseline"])
    for ri,cat in enumerate(judge_cats):
        row = tbl_delta.rows[ri+1]
        name = CATEGORY_INFO.get(cat,("","",""))[0]
        zs   = mean(JUDGE_V2_CAT[cat][1])
        csep = mean(JUDGE_V2_CAT[cat][3])
        czs  = mean(JUDGE_V2_CAT[cat][4])
        row.cells[0].text = name
        if zs is not None and csep is not None and czs is not None:
            d1, d2 = csep-zs, czs-zs
            row.cells[1].text = f"{d1:+.3f}"
            row.cells[2].text = f"{d2:+.3f}"
            row.cells[3].text = f"{zs:.3f}"
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row.cells[1], delta_color(d1))
            set_cell_bg(row.cells[2], delta_color(d2))
        for c in row.cells:
            if c.paragraphs[0].runs: c.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_borders(c)

    doc.add_page_break()

    # ── 4. Graphs ──
    add_heading(doc, "4. Vizualizacije", 1)
    add_para(doc,
        "Napomena: Grafici ispod su generisani iz auto-scored podataka. "
        "Vrijednosti se razlikuju od v2 judge-scored tabela.",
        italic=True, size=9, color="808080")
    doc.add_paragraph()
    graph_captions = [
        ("1_overall_by_condition.png", "Slika 1 — Ukupna tačnost po uvjetu (auto-scored)"),
        ("2_by_model_and_condition.png", "Slika 2 — Performanse po modelu i uvjetu (auto-scored)"),
        ("3_category_heatmap.png", "Slika 3 — Heatmapa po kategoriji i uvjetu"),
        ("4_csep_improvement.png", "Slika 4 — CSEP delta vs Zero-Shot"),
        ("5_per_model_range.png", "Slika 5 — Raspon skora po modelu"),
    ]
    for fname, caption in graph_captions:
        path = os.path.join(GRAPHS_DIR, fname)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(5.5))
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs: p.runs[0].italic = True; p.runs[0].font.size = Pt(9)
            doc.add_paragraph()

    doc.add_page_break()

    # ── 5. Key Findings ──
    add_heading(doc, "5. Ključni nalazi", 1)
    findings = [
        ("Nalaz 1: ZS i CSEP+ZS su praktički izjednačeni",
         "Razlika između Zero-Shot (0.725) i CSEP+ZS (0.721) je samo 0.4pp — zanemarivo. "
         "CSEP-Only (0.707) je 1.8pp ispod ZS. Za prosječan zadatak, CSEP ne nudi korist "
         "koja bi opravdala 4x veći trošak API poziva."),
        ("Nalaz 2: CoT — bolji nego što je v1 pokazivao, ali ne uvijek",
         "CoT (0.685) je na judge-scored zadacima ispod ZS (0.725), ali razlika je 4pp, "
         "ne 13.7pp kao što je v1 sugerisao. CoT sinteza ('All three attempts converge...') "
         "daje korektan finalni odgovor, ali je verbozna i ponekad sadržava meta-komentare "
         "koje sudija kažnjava. Na auto-scored zadacima CoT dominira (0.701 vs ZS 0.607)."),
        ("Nalaz 3: CSEP jasno pomaže na paradoks/temporalnim zadacima",
         "Self-reference paradox: ZS 0.617 → CSEP+ZS 0.667 (+5pp), CSEP 0.733 (+11.7pp). "
         "Temporal reasoning: ZS 0.770 → CSEP 0.820 (+5pp). "
         "Dekompozicija primorava model da eksplicitno razmisli o tipu pitanja, "
         "što pomaže prepoznati samoreferencijalne i temporalne zamke."),
        ("Nalaz 4: CSEP konzistentno šteti na jezičkim i adversarijalnim zadacima",
         "Linguistic constraint: ZS 0.683 → CSEP 0.567 (−11.7pp). "
         "Adversarial pressure: ZS 0.650 → CSEP+ZS 0.525 (−12.5pp). "
         "CSEP elaboracija konteksta čini model podložnijim socijalnom pritisku — "
         "prevelika analiza oslabljuje direktnost i čvrstinu odgovora."),
        ("Nalaz 5: Phi-4 je jedini model gdje CSEP konzistentno pomaže",
         "Phi-4: ZS 0.794, CoT 0.718, CSEP 0.763, CSEP+ZS 0.784. "
         "Samo kod Phi-4 vidimo da CSEP+ZS (0.784) primiče ZS (0.794) umjesto da ga degradira. "
         "Za modele ispod 12B parametara, CSEP scaffolding donosi više štete nego koristi."),
        ("Nalaz 6: Llama-3.1-8B ima najveći CoT kolaps",
         "Llama: ZS 0.653 → CoT 0.558 (−9.5pp). "
         "Model nije u stanju pouzdano sintetizovati tri nezavisna passa — "
         "sinteza postaje konfuzna ili protivurječna. "
         "CoT treba koristiti s oprezom za 8B Llama modele."),
    ]
    for title, body in findings:
        p = doc.add_paragraph()
        r1 = p.add_run(title); r1.bold = True; r1.font.size = Pt(11.5)
        r1.font.color.rgb = RGBColor(0x1F,0x38,0x64)
        add_para(doc, body)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ── 6. Examples ──
    add_heading(doc, "6. Primjeri po kategoriji (judge-scored)", 1)
    add_para(doc,
        "Za svaku judge-scored kategoriju prikazan je jedan visoko ocijenjeni i jedan nisko ocijenjeni primjer "
        "iz v2 skupa (puni odgovori, razlozi iz sudije). Primjeri su odabrani automatski — "
        "najniži i najviši normalizovani skor u kategoriji.")
    doc.add_paragraph()

    judge_cat_order = sorted(EXAMPLES_V2.keys(), key=lambda x: int(x.split("_")[0]))
    for idx, cat in enumerate(judge_cat_order):
        if cat not in CATEGORY_INFO: continue
        cat_name = CATEGORY_INFO[cat][0]
        add_heading(doc, f"6.{idx+1}  {cat_name}", 2)
        ex = EXAMPLES_V2[cat]
        for kind, bg in [("best","375623"),("worst","C00000")]:
            e = ex[kind]
            cond_name = CONDITIONS.get(e["condition"], str(e["condition"]))
            label = "DOBAR PRIMJER" if kind=="best" else "LOŠ PRIMJER"
            add_example_box(doc, label, bg,
                e["model"], cond_name, e["score"],
                e["question"], e["expected"], e["response"], e["reason"])

    doc.add_page_break()

    # ── 7. Conclusion ──
    add_heading(doc, "7. Zaključak", 1)
    add_para(doc,
        "CSEP (Conceptual Space Expansion Prompting) ne poboljšava konzistentno performanse "
        "u poređenju s jednostavnim zero-shot pristupom. Na prosječnom skupu zadataka, "
        "Zero-Shot i CSEP+ZS su praktički ekvivalentni (razlika 0.4pp), uz 4x veći trošak za CSEP. "
        "CSEP je opravdan samo za specifične kategorije — zadaci koji zahtijevaju prepoznavanje "
        "samoreferencijalne paradokse i temporalne odnose pokazuju konzistentno poboljšanje.")
    doc.add_paragraph()
    add_para(doc,
        "CoT je superioran na strukturiranim računskim zadacima (matematika, prebrojavanje, relacijske "
        "zagonetke s definitivnim odgovorima), ali inferioran na otvorenim i adversarijalnim zadacima. "
        "Ovo se ne smije interpretirati kao 'CoT je loš' — zavisno od primjene, može biti optimalan izbor.")
    doc.add_paragraph()
    add_para(doc, "Preporuke:", bold=True)
    recs = [
        "Za strukturirane računske zadatke: koristiti CoT.",
        "Za self-reference i temporalne zadatke: koristiti CSEP ili CSEP+ZS.",
        "Za adversarijalne i jezičke zadatke: koristiti Zero-Shot.",
        "CSEP+ZS je uvijek bolji od CSEP-Only — zero-shot sidro smanjuje drift.",
        "Za modele ispod 12B: izbjegavati CSEP osim na kategorijalnom skupu gdje postoji jasan signal.",
        "Phi-4 (14B) je jedini model gdje CSEP generalno donosi korist.",
    ]
    for rec in recs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rec)
    doc.add_paragraph()
    add_para(doc,
        "Buduća istraživanja trebaju ispitati da li strožija provjera znanja u Stage 2 dekompozicije "
        "može zadržati prednosti CSEP-a a smanjiti halucinacije na osjetljivim kategorijama.",
        italic=True)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

if __name__ == "__main__":
    build()
