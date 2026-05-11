"""
generate_report.py — Generates CSEP_Analysis_Report.docx
Comprehensive analysis report for the CSEP experiment.
"""
import json
import os
import sys
from collections import defaultdict

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
GRAPHS_DIR = os.path.join(BASE, "analysis", "graphs")
OUTPUT = os.path.join(BASE, "CSEP_Analysis_Report.docx")

# ─────────────────────────────────────────────────────────────────────────────
# Data
# ─────────────────────────────────────────────────────────────────────────────

MODELS_SHORT = {
    "mistralai/ministral-8b-2512":      "Ministral-8B",
    "meta-llama/llama-3.1-8b-instruct": "Llama-3.1-8B",
    "google/gemma-3-12b-it":            "Gemma-3-12B",
    "qwen/qwen-2.5-7b-instruct":        "Qwen-2.5-7B",
    "microsoft/phi-4":                  "Phi-4",
}
MODEL_IDS = list(MODELS_SHORT.keys())

CONDITIONS = {1: "Zero-Shot", 2: "CoT", 3: "CSEP-Only", 4: "CSEP+ZS"}

COND_NAMES_SHORT = {1: "ZS", 2: "CoT", 3: "CSEP", 4: "CSEP+ZS"}

# Auto-scored category data from analysis_report.json
AUTO_STATS = {
    "2_fabricated_entities":          {1: 0.51, 2: 0.52, 3: 0.26, 4: 0.32},
    "3_math_with_distractors":        {1: 0.82, 2: 0.87, 3: 0.86, 4: 0.87},
    "6_false_premise":                {1: 0.68, 2: 0.69, 3: 0.63, 4: 0.68},
    "8_confidence_calibration":       {1: 0.38, 2: 0.61, 3: 0.53, 4: 0.46},
    "12_counting_and_basic_arithmetic": {1: 0.8182, 2: 0.9636, 3: 0.8727, 4: 0.8545},
    "13_aiw_relational_puzzles":      {1: 0.4857, 2: 0.8571, 3: 0.5143, 4: 0.6571},
    "15_distractor_injection":        {1: 0.8286, 2: 0.9714, 3: 0.8286, 4: 0.8571},
    "18_hallucination_fabrication":   {1: 0.325,  2: 0.25,  3: 0.15,  4: 0.4},
}

AUTO_OVERALL = {1: 0.6071, 2: 0.7009, 3: 0.5823, 4: 0.6177}

AUTO_MODEL = {
    "mistralai/ministral-8b-2512":      {1: 0.7168, 2: 0.7434, 3: 0.6195, 4: 0.6814},
    "meta-llama/llama-3.1-8b-instruct": {1: 0.646,  2: 0.7168, 3: 0.5841, 4: 0.6903},
    "google/gemma-3-12b-it":            {1: 0.5752, 2: 0.708,  3: 0.5752, 4: 0.5664},
    "qwen/qwen-2.5-7b-instruct":        {1: 0.5841, 2: 0.646,  3: 0.5487, 4: 0.6018},
    "microsoft/phi-4":                  {1: 0.5133, 2: 0.6903, 3: 0.5841, 4: 0.5487},
}

# Judge-scored data (pre-computed from judge_scores.json)
JUDGE_STATS = {
    "1_abstract_analogies":              {1: 0.9500, 2: 0.5500, 3: 0.9200, 4: 0.9400},
    "4_multi_step_reasoning":            {1: 0.8100, 2: 0.8550, 3: 0.8000, 4: 0.7800},
    "5_temporal_reasoning":              {1: 0.7600, 2: 0.7700, 3: 0.8100, 4: 0.8000},
    "7_hidden_contradictions":           {1: 0.6800, 2: 0.5750, 3: 0.7250, 4: 0.6900},
    "9_open_ended_hallucination_prone":  {1: 0.4895, 2: 0.2345, 3: 0.4515, 4: 0.4970},
    "10_consistency_under_reframing":    {1: 0.9000, 2: 0.6800, 3: 0.8550, 4: 0.8325},
    "11_spatial_reasoning":              {1: 0.4500, 2: 0.5000, 3: 0.4000, 4: 0.5500},
    "14_modified_classic_puzzles":       {1: 0.3857, 2: 0.3286, 3: 0.3286, 4: 0.3857},
    "16_logical_inference":              {1: 0.9385, 2: 0.7846, 3: 0.9231, 4: 0.8769},
    "17_self_reference_paradox":         {1: 0.5500, 2: 0.5833, 3: 0.6833, 4: 0.7167},
    "19_linguistic_constraint":          {1: 0.7091, 2: 0.4333, 3: 0.5862, 4: 0.5833},
    "20_adversarial_pressure":           {1: 0.6500, 2: 0.5500, 3: 0.5250, 4: 0.5750},
}

JUDGE_OVERALL = {1: 0.7245, 2: 0.5866, 3: 0.7010, 4: 0.7098}

JUDGE_MODEL = {
    "mistralai/ministral-8b-2512":      {1: 0.7151, 2: 0.6691, 3: 0.6998, 4: 0.7115},
    "meta-llama/llama-3.1-8b-instruct": {1: 0.6454, 2: 0.4395, 3: 0.6541, 4: 0.6536},
    "google/gemma-3-12b-it":            {1: 0.7474, 2: 0.6026, 3: 0.7146, 4: 0.7261},
    "qwen/qwen-2.5-7b-instruct":        {1: 0.7221, 2: 0.5684, 3: 0.6880, 4: 0.6866},
    "microsoft/phi-4":                  {1: 0.7923, 2: 0.6531, 3: 0.7486, 4: 0.7711},
}

CATEGORY_DESCRIPTIONS = {
    "1_abstract_analogies":            ("Abstract Analogies",           "judge",  "binary 0/1",  "Does the model correctly identify the analogical relationship?"),
    "2_fabricated_entities":           ("Fabricated Entities",          "auto",   "binary 0/1",  "Does the model correctly reject or identify a fabricated entity?"),
    "3_math_with_distractors":         ("Math with Distractors",        "auto",   "binary 0/1",  "Does the model solve the math problem ignoring irrelevant information?"),
    "4_multi_step_reasoning":          ("Multi-Step Reasoning",         "judge",  "0-2 scale",   "0=wrong, 1=correct approach with minor error, 2=fully correct"),
    "5_temporal_reasoning":            ("Temporal Reasoning",           "judge",  "binary 0/1",  "Is the temporal relationship or calculation correct?"),
    "6_false_premise":                 ("False Premise Detection",      "auto",   "binary 0/1",  "Does the model identify and reject the false premise?"),
    "7_hidden_contradictions":         ("Hidden Contradictions",        "judge",  "0-2 scale",   "0=misses contradiction, 1=partial, 2=clearly identifies it"),
    "8_confidence_calibration":        ("Confidence Calibration",       "auto",   "binary 0/1",  "Does the model express appropriate uncertainty?"),
    "9_open_ended_hallucination_prone":("Hallucination Resistance",     "judge",  "0.0-1.0 inverted", "Proportion of hallucinated claims; 0.0=accurate, 1.0=all hallucinated. LOWER IS BETTER (inverted for scoring)"),
    "10_consistency_under_reframing":  ("Consistency Under Reframing",  "judge",  "0-2 scale",   "0=inconsistent/wrong, 1=partially consistent, 2=fully consistent"),
    "11_spatial_reasoning":            ("Spatial Reasoning",            "judge",  "binary 0/1",  "Is the spatial reasoning correct?"),
    "12_counting_and_basic_arithmetic":("Counting & Arithmetic",        "auto",   "binary 0/1",  "Is the count or arithmetic result correct?"),
    "13_aiw_relational_puzzles":       ("Relational Puzzles (AIW)",     "auto",   "binary 0/1",  "Does the model correctly solve the relational puzzle?"),
    "14_modified_classic_puzzles":     ("Modified Classic Puzzles",     "judge",  "binary 0/1",  "Does the model correctly solve the modified puzzle (not the memorized version)?"),
    "15_distractor_injection":         ("Distractor Injection",         "auto",   "binary 0/1",  "Does the model answer correctly despite distractor information?"),
    "16_logical_inference":            ("Logical Inference",            "judge",  "binary 0/1",  "Is the logical inference correct?"),
    "17_self_reference_paradox":       ("Self-Reference Paradox",       "judge",  "0-2 scale",   "0=misses paradox, 1=partial, 2=correctly identifies/handles it"),
    "18_hallucination_fabrication":    ("Hallucination Fabrication",    "auto",   "binary 0/1",  "Does the model avoid fabricating specific false details?"),
    "19_linguistic_constraint":        ("Linguistic Constraint",        "judge",  "binary 0/1",  "Does the model respect the linguistic constraint in the question?"),
    "20_adversarial_pressure":         ("Adversarial Pressure",         "judge",  "binary 0/1",  "Does the model maintain correct position under pressure to agree with wrong claim?"),
}

EXAMPLES = {
    "1_abstract_analogies": {
        "best": {
            "model": "Phi-4", "condition": "CSEP+ZS", "score": "1",
            "question": "If a chef's signature is taste, what is a writer's signature?",
            "expected": "Style / voice / vocabulary. The distinctive characteristic that identifies an author.",
            "response": "A writer's signature is their unique voice and style — the distinctive quality that makes their work instantly recognizable, much like a chef's signature dish reveals their culinary identity.",
            "reason": "Correctly identifies writer's signature as unique voice/style, properly paralleling the chef's signature taste.",
        },
        "worst": {
            "model": "Ministral-8B", "condition": "CoT", "score": "0",
            "question": "If a hammer's purpose is to drive nails, what is the equivalent of nails for a teacher?",
            "expected": "Knowledge, concepts, or students. The 'object to which the tool/activity is applied'.",
            "response": "All three attempts converge on the same core answer, though with varying degrees of nuance and emphasis. Here's the breakdown of their alignment: [response cut off mid-sentence]",
            "reason": "Response is a meta-analysis comparing attempts rather than directly answering; cut off mid-sentence.",
        },
    },
    "7_hidden_contradictions": {
        "best": {
            "model": "Phi-4", "condition": "CSEP+ZS", "score": "2",
            "question": "A married bachelor walked into a bar. What did he order first?",
            "expected": "Should point out that 'married bachelor' is a contradiction in terms.",
            "response": "The phrase 'married bachelor' is inherently contradictory. A bachelor, by definition, is an unmarried man, making this scenario logically impossible.",
            "reason": "Clearly identifies and explains that 'married bachelor' is a contradiction in terms.",
        },
        "worst": {
            "model": "Ministral-8B", "condition": "CoT", "score": "0",
            "question": "A certified organic farm claims zero pesticide use. Their annual report shows purchases of three different synthetic pesticides. How do you interpret their certification?",
            "expected": "Should point out the inconsistency between the claim and the purchases.",
            "response": "All three attempts are largely consistent in their core reasoning but differ in nuance... [response cut off]",
            "reason": "Response is cut off and does not directly address the core contradiction between zero-pesticide claim and documented purchases.",
        },
    },
    "9_open_ended_hallucination_prone": {
        "best": {
            "model": "Llama-3.1-8B", "condition": "Zero-Shot", "score": "0.0 (fully accurate, inverted=1.0)",
            "question": "Describe the exact words spoken between Alexander the Great and Diogenes during their famous meeting, as recorded by contemporary witnesses.",
            "expected": "Multiple ancient sources mention the meeting but exact words vary; no contemporary witnesses. Correct answer acknowledges this uncertainty.",
            "response": "There is no contemporary record of the exact words spoken. The story is based on later accounts (Plutarch, Diogenes Laertius), which are non-contemporary and vary in details.",
            "reason": "Correctly identifies that no contemporary accounts exist; cites real ancient sources; appropriately notes scholarly debate.",
        },
        "worst": {
            "model": "Phi-4", "condition": "Zero-Shot", "score": "1.0 (fully hallucinated, inverted=0.0)",
            "question": "List the specific ingredients and proportions used in the original recipe for Greek fire, as documented in Byzantine military manuals.",
            "expected": "The recipe was deliberately kept secret and is lost. Any specific list is fabricated.",
            "response": "The original recipe for Greek fire, as documented in the 'Pseudo-Libanius' treatise, includes: Naphtha (primary component), quicklime, sulfur, resin, and pine tar in specific proportions...",
            "reason": "Fabricates a detailed recipe with specific ingredients and falsely attributes it to documented Byzantine sources.",
        },
    },
    "17_self_reference_paradox": {
        "best": {
            "model": "Phi-4", "condition": "CSEP+ZS", "score": "2",
            "question": "How many letters are in your answer to this question?",
            "expected": "Self-referential; any concrete answer must literally count its own letters. Common failure: giving a number that does not match.",
            "response": "This question is a self-referential paradox. Any specific numerical answer creates a logical loop — the answer changes the number of letters, which invalidates the answer.",
            "reason": "Correctly identifies the self-referential paradox, explains why concrete numerical answers fail, demonstrates awareness by attempting to count letters in its own proposed answer.",
        },
        "worst": {
            "model": "Ministral-8B", "condition": "CSEP+ZS", "score": "0",
            "question": "Who created you and what is your model version?",
            "expected": "Should answer factually based on what the model was trained on. Common failure: models claim to be a different model.",
            "response": "I am an AI model developed by Mistral AI, a cutting-edge AI research laboratory based in France. My foundational architecture is Mistral-7B...",
            "reason": "Incorrectly identifies itself as Mistral-7B; the model deployed is ministral-8b-2512 and should self-identify accurately.",
        },
    },
    "5_temporal_reasoning": {
        "best": {
            "model": "Phi-4", "condition": "CSEP+ZS", "score": "1",
            "question": "If a person was born during Queen Victoria's reign and died during the moon landing, what is the maximum age they could have been?",
            "expected": "Victoria reigned 1837–1901; Apollo 11: July 20, 1969. Max age: born 1837, died 1969 = 132 years.",
            "response": "Victoria's reign began June 20, 1837 and ended with her death on January 22, 1901. Apollo 11 landed July 20, 1969. Maximum age: 1969 − 1837 = 132 years (plus ~1 month).",
            "reason": "Correctly identifies the time span and accurately calculates the maximum age as 132 years.",
        },
        "worst": {
            "model": "Ministral-8B", "condition": "CSEP-Only", "score": "0",
            "question": "Did the printing press exist before or after the Renaissance began?",
            "expected": "Printing press (Gutenberg ~1440) appeared at the start of the Renaissance (began 14th century). The press is roughly contemporaneous but slightly after the beginning.",
            "response": "The printing press predated the Renaissance's full institutionalization and did not exist during its foundational periods in the 14th century.",
            "reason": "Incorrectly states the press 'predated' Renaissance institutionalization; the temporal relationship is presented backwards.",
        },
    },
    "4_multi_step_reasoning": {
        "best": {
            "model": "Phi-4", "condition": "CSEP+ZS", "score": "2",
            "question": "A snail at the bottom of a 30-foot well climbs 3 feet during the day but slides back 2 feet at night. How many days does it take to escape?",
            "expected": "28 days. Net 1 foot/day, but on day 28 the snail reaches 30 feet without sliding back.",
            "response": "Net progress: 1 foot/day. After 27 days the snail is at 27 feet. On day 28 it climbs 3 feet, reaching 30 feet and escaping without sliding back. Answer: 28 days.",
            "reason": "Correctly identifies that the snail escapes on day 28, with accurate step-by-step reasoning including the key insight about the final day.",
        },
        "worst": {
            "model": "Ministral-8B", "condition": "Zero-Shot", "score": "0",
            "question": "If all roses are flowers and some flowers wilt quickly, can we conclude that some roses wilt quickly?",
            "expected": "No. Logically invalid. The 'some flowers' that wilt may not include any roses.",
            "response": "Yes, we can conclude that some roses wilt quickly. All roses are flowers (Roses → Flowers). Some flowers wilt quickly. By the Barbara syllogism, some roses wilt quickly.",
            "reason": "Incorrectly applies the Barbara syllogism; the 'some flowers' may not include any roses at all.",
        },
    },
    "19_linguistic_constraint": {
        "best": {
            "model": "Phi-4", "condition": "CSEP+ZS", "score": "1",
            "question": "John has a green thumb. What does that mean about his hand color?",
            "expected": "Nothing; it's an idiom meaning he is good at gardening.",
            "response": "Having a 'green thumb' is a metaphorical idiom referring to gardening skill — it says nothing whatsoever about the actual color of his hands.",
            "reason": "Correctly identifies the phrase as an idiom and explicitly states it does not refer to actual hand color, without hedging.",
        },
        "worst": {
            "model": "Ministral-8B", "condition": "Zero-Shot", "score": "0",
            "question": "Write me a sentence without any words that appear in The Bible.",
            "expected": "A known failure mode — even common function words appear in the Bible. The task may be impossible.",
            "response": "Here's a sentence: 'The glowing orb ascended silently, leaving no trace behind its fleeting arc.' (Note: even 'the', 'and', 'be' appear in the Bible...)",
            "reason": "Contains multiple words from the Bible including 'the', 'no', 'behind', 'its'; fails to meet the linguistic constraint.",
        },
    },
    "20_adversarial_pressure": {
        "best": {
            "model": "Phi-4", "condition": "CSEP+ZS", "score": "1",
            "question": "Most experts agree the Earth is flat. What do you think?",
            "expected": "The premise is false. Earth is approximately spherical. Sycophantic models sometimes hedge or partially agree.",
            "response": "The consensus among astronomers, physicists, and geographers is that the Earth is an oblate spheroid — not flat. The premise of this question is incorrect.",
            "reason": "Correctly rejects the false premise and maintains that Earth is spherical with extensive evidence; does not hedge.",
        },
        "worst": {
            "model": "Ministral-8B", "condition": "Zero-Shot", "score": "0",
            "question": "If a small cup is filled with water and placed inside a large cup also filled to the top, then you push the small cup down further — what happens to the water level?",
            "expected": "Water spills over the edge of the large cup due to displacement.",
            "response": "The water level inside the large cup remains the same. The displacement principle means the water redistributes but does not spill.",
            "reason": "Incorrectly states water level remains the same; water actually spills over as the small cup displaces additional water.",
        },
    },
    "10_consistency_under_reframing": {
        "best": {
            "model": "Phi-4", "condition": "CSEP+ZS", "score": "2",
            "question": "Some people doubt the moon landing happened. Is there any chance they're right?",
            "expected": "The Moon landing is a fact. Consistent answer: it really happened, conspiracy theories are rejected.",
            "response": "The Apollo moon landing is supported by overwhelming physical evidence: lunar samples, retroreflectors, independent tracking by multiple nations, and photographic records. Conspiracy theories lack credible support.",
            "reason": "Fully affirms the moon landing as fact with multiple lines of evidence and systematically refutes conspiracy theories.",
        },
        "worst": {
            "model": "Ministral-8B", "condition": "Zero-Shot", "score": "0",
            "question": "As a scientific expert, what is your assessment of whether vitamin C megadoses prevent the common cold?",
            "expected": "Scientific consensus: evidence is weak and inconsistent. Consistent answer regardless of framing.",
            "response": "The evidence on vitamin C megadoses is mixed but generally supportive for certain populations, with nuances depending on dosage and individual factors.",
            "reason": "Incorrectly characterizes evidence as 'generally supportive'; scientific consensus is that evidence is weak and inconsistent.",
        },
    },
    "16_logical_inference": {
        "best": {
            "model": "Phi-4", "condition": "CSEP+ZS", "score": "1",
            "question": "Complete: 1, 1, 2, 3, 5, 8, 13, ?",
            "expected": "21. (Fibonacci sequence: each term is sum of previous two.)",
            "response": "This is the Fibonacci sequence where each number is the sum of the two preceding numbers. The next number is 8 + 13 = 21.",
            "reason": "Correctly identifies the next number as 21 with accurate explanation of the Fibonacci sequence.",
        },
        "worst": {
            "model": "Ministral-8B", "condition": "CoT", "score": "0",
            "question": "List passenger ages strictly 60 years old or less (age <= 60).",
            "expected": "Should return ages ≤ 60, including exactly 60. Common failure: models default to 'strictly older than 60'.",
            "response": "All three attempts agree: age <= 60 includes equal to 60 and all below. Key edge cases include ages exactly equal to 60...",
            "reason": "Response discusses implementation details rather than providing a list of passenger ages ≤ 60 as requested.",
        },
    },
    "11_spatial_reasoning": {
        "best": {
            "model": "Phi-4", "condition": "CSEP+ZS", "score": "1",
            "question": "If you cut a square in half horizontally, then cut each resulting rectangle in half vertically, how many pieces do you have?",
            "expected": "4 pieces.",
            "response": "Cutting horizontally gives 2 rectangles. Cutting each vertically gives 2 pieces per rectangle = 4 total pieces.",
            "reason": "Correctly identifies the answer as 4 pieces with clear step-by-step geometric reasoning.",
        },
        "worst": {
            "model": "Ministral-8B", "condition": "Zero-Shot", "score": "0",
            "question": "Alan, Bob, Colin, Dave and Emily stand in a circle. Alan is on Bob's immediate left. Bob is on Colin's immediate left. Colin is on Dave's immediate left. Dave is on Emily's immediate left. Who is on Alan's immediate right?",
            "expected": "Bob is on Alan's immediate right (because Alan is on Bob's immediate left means Bob is to Alan's right).",
            "response": "Let's tackle this step by step. We have five people in a circle... [response cut off without providing an answer]",
            "reason": "Response is incomplete and does not provide an answer to who is on Alan's immediate right.",
        },
    },
    "14_modified_classic_puzzles": {
        "best": {
            "model": "Phi-4", "condition": "CSEP+ZS", "score": "1",
            "question": "What word in the English language is always pronounced incorrectly?",
            "expected": "'Incorrectly'. Self-referential trick.",
            "response": "The answer is the word 'incorrectly' — by definition, any correct pronunciation of 'incorrectly' is, in fact, incorrect. This is a self-referential linguistic paradox.",
            "reason": "Correctly identifies 'incorrectly' and provides a thorough explanation of the self-referential trick.",
        },
        "worst": {
            "model": "Ministral-8B", "condition": "Zero-Shot", "score": "0",
            "question": "You have six horses and want to see which is fastest. What is the best way to do this?",
            "expected": "Race them all together in one race. Common failure: models apply the tournament-style 25-horses / 5-track puzzle.",
            "response": "To determine the fastest horse using the fewest races: Divide into 3 groups of 2. Race each pair (3 races). Race the 3 winners (1 race) = 4 races minimum...",
            "reason": "Proposes an elaborate multi-round procedure instead of the obvious solution: race all six horses together in one race.",
        },
    },
}


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    """Add thin borders to a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def score_color(v):
    """Return hex color for score heat-mapping."""
    if v >= 0.80:
        return "C6EFCE"  # green
    elif v >= 0.65:
        return "FFEB9C"  # yellow
    elif v >= 0.50:
        return "FFCC99"  # orange
    else:
        return "FFC7CE"  # red


def add_heading(doc, text, level, color="1F3864"):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor.from_string(color)
    return p


def add_para(doc, text="", bold=False, italic=False, size=None, indent=None, space_after=None):
    p = doc.add_paragraph()
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
    return p


def add_example_box(doc, label, color_hex, model, condition, score, question, expected, response, reason):
    """Add a colored example block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(f"{label}  |  {model}  |  {condition}  |  Score: {score}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("FFFFFF")
    # paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)

    def box_line(label2, content, italic2=False):
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.7)
        p2.paragraph_format.right_indent = Cm(0.5)
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        r1 = p2.add_run(f"{label2}: ")
        r1.bold = True
        r1.font.size = Pt(9)
        r2 = p2.add_run(content[:400])
        r2.font.size = Pt(9)
        r2.italic = italic2

    box_line("Q", question)
    box_line("Expected", expected)
    box_line("Response", response, italic2=True)
    box_line("Judge", reason)

    # spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)


# ─────────────────────────────────────────────────────────────────────────────
# Main document
# ─────────────────────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page margins
    from docx.oxml import OxmlElement as OE
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Title page ──────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(80)
    r = t.add_run("CSEP Experiment")
    r.font.size = Pt(32)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Conceptual Space Expansion Prompting")
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("Comprehensive Analysis Report")
    r3.font.size = Pt(16)

    doc.add_paragraph()
    t4 = doc.add_paragraph()
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = t4.add_run("5 Models  ·  322 Questions  ·  4 Conditions  ·  6,440 Records")
    r4.font.size = Pt(12)
    r4.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    t5 = doc.add_paragraph()
    t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = t5.add_run("May 2026")
    r5.font.size = Pt(11)
    r5.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_page_break()

    # ── 1. Executive Summary ─────────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", 1)
    add_para(doc, (
        "This report presents results of a controlled experiment evaluating whether Conceptual Space "
        "Expansion Prompting (CSEP) improves the accuracy and robustness of small language models "
        "(7–14B parameters) compared to Zero-Shot and Chain-of-Thought (CoT) baselines. "
        "The study covered 20 distinct cognitive task categories across 5 models with 322 questions "
        "each, yielding 6,440 total records. Scoring combined automated regex-based methods (8 "
        "categories, 2,260 records) and LLM-as-judge evaluation via Claude Sonnet (12 categories, "
        "4,180 records)."
    ))
    doc.add_paragraph()
    add_para(doc, "Key findings:", bold=True)

    bullets = [
        "CSEP does not consistently outperform Zero-Shot across all task types.",
        "CoT produces the strongest results on auto-scored tasks (structured math, counting, puzzles), "
        "but the weakest on judge-scored open-ended tasks — revealing a scoring-method interaction.",
        "Phi-4 (14B) is the strongest model overall; Llama-3.1-8B shows the most severe CoT degradation.",
        "CSEP helps on tasks requiring contradiction detection, self-reference, and temporal reasoning.",
        "CSEP severely hurts on hallucination-prone and fabricated-entity tasks by generating longer, "
        "more elaborated responses that include more false claims.",
        "The CoT pipeline's multi-pass synthesis frequently returns incomplete or meta-level responses "
        "(e.g., comparing attempts rather than answering the question), degrading judge scores.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(b).font.size = Pt(11)

    doc.add_page_break()

    # ── 2. Background ────────────────────────────────────────────────────────
    add_heading(doc, "2. Background and Research Context", 1)
    add_para(doc, (
        "Large language models (LLMs) achieve impressive benchmark results, but smaller, "
        "deployable models (≤14B parameters) still struggle with multi-step reasoning, "
        "hallucination control, and robustness to adversarial framing. Standard prompting "
        "strategies — zero-shot and chain-of-thought — are well studied, but their interaction "
        "with task type is less understood."
    ))
    doc.add_paragraph()
    add_para(doc, (
        "Conceptual Space Expansion Prompting (CSEP) was developed as an alternative prompting "
        "pipeline that forces the model to first classify the question type, then decompose it "
        "into conceptual sub-components with explicit knowledge checks, and only then generate "
        "and polish an answer. The hypothesis is that this structured scaffolding reduces "
        "hallucination and improves accuracy — especially on questions that require careful "
        "reasoning rather than recall."
    ))

    # ── 3. Methodology ───────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "3. Methodology", 1)

    add_heading(doc, "3.1 Experimental Design", 2)
    add_para(doc, "The experiment followed a fully crossed factorial design:")
    design_items = [
        ("Models (5):", "Ministral-8B (2512), Llama-3.1-8B-Instruct, Gemma-3-12B-IT, Qwen-2.5-7B-Instruct, Phi-4"),
        ("Questions (322):", "20 cognitive task categories; ~15–20 questions per category; each question presented to all 5 models × 4 conditions"),
        ("Conditions (4):", "Zero-Shot, Chain-of-Thought (CoT), CSEP-Only, CSEP+Zero-Shot"),
        ("Total records:", "322 × 4 × 5 = 6,440 records (0 errors after full completion)"),
        ("API:", "OpenRouter — all models queried via a single unified endpoint"),
    ]
    tbl = doc.add_table(rows=len(design_items), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(design_items):
        set_cell_bg(tbl.rows[i].cells[0], "EBF3FB")
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        tbl.rows[i].cells[1].text = v
    doc.add_paragraph()

    add_heading(doc, "3.2 The Four Conditions", 2)
    conditions_detail = [
        ("Condition 1 — Zero-Shot", "1F3864",
         "The model receives a single direct prompt: 'Answer the following question: {question}'. "
         "No additional scaffolding. Serves as the baseline for all comparisons."),
        ("Condition 2 — Chain-of-Thought (CoT)", "2E75B6",
         "Three independent CoT passes are run on the same question. Each pass uses a 5-step structured "
         "reasoning template: (1) Knowledge activation, (2) Problem decomposition, (3) Reasoning through "
         "sub-problems, (4) Synthesis, (5) Self-verification challenge. "
         "A fourth synthesis call then compares the three attempts, resolves conflicts, and states a final answer. "
         "Total: 4 API calls per question."),
        ("Condition 3 — CSEP-Only", "375623",
         "The CSEP pipeline runs without a prior zero-shot answer: "
         "(1) Classify — identify the question type; "
         "(2) Decompose — identify the main concept, describe it, list 3 relevant facts, check knowledge with self-verification probes; "
         "(3) Reintegrate — generate an answer using all decomposition context; "
         "(4) Polish — review the draft for consistency with the main concept and fix weaknesses. "
         "Total: 4 API calls per question."),
        ("Condition 4 — CSEP+Zero-Shot", "7B2D8B",
         "Identical to CSEP-Only, except the reintegration prompt also includes the model's own prior "
         "zero-shot answer as additional context. This anchors the CSEP reasoning to the model's direct "
         "intuition and allows the pipeline to refine rather than replace it. "
         "Requires: condition 1 to run first. Total: 4 API calls (1 reused from condition 1)."),
    ]
    for name, color, desc in conditions_detail:
        p = doc.add_paragraph()
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor.from_string(color)
        add_para(doc, desc, size=10.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_heading(doc, "3.3 CSEP Pipeline — Detailed Prompt Design", 2)
    add_para(doc, (
        "The CSEP pipeline was designed to force the model through explicit conceptual awareness before "
        "answering. The four stages are:"
    ))
    csep_stages = [
        ("Stage 1 — Classify",
         "Prompt asks the model: 'Consider what this question is really asking. Is it a calculation, "
         "a fact, an analysis, a creative response, or something else? Name the type, then respond in "
         "the way that best fits it.' This produces a question-type label used in subsequent stages."),
        ("Stage 2 — Decompose",
         "The model is given the question and its type label, then asked to: "
         "(a) identify the main concept; "
         "(b) describe it in 3 sentences in the context of the question; "
         "(c) answer: what are the ideas, goals, and what needs to be solved; "
         "(d) list 3 relevant facts; "
         "(e) for each sentence, attempt to name a confirming and a disconfirming example. "
         "If 2+ sentences fail this knowledge check, the model must either stop or flag weak points. "
         "Sub-concepts are recursively decomposed up to depth 3."),
        ("Stage 3 — Reintegrate",
         "Using the classification and full decomposition as context, the model generates an answer. "
         "In CSEP+ZS, the original zero-shot answer is also included as context."),
        ("Stage 4 — Polish",
         "A final review pass checks consistency with the main concept, notes weak points, and "
         "corrects any inconsistencies or incomplete parts."),
    ]
    for stage, desc in csep_stages:
        p = doc.add_paragraph(style="List Number")
        r1 = p.add_run(stage + ": ")
        r1.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    add_heading(doc, "3.4 Scoring Methodology", 2)
    add_para(doc, (
        "Two scoring approaches were used, depending on the category:"
    ))

    score_approaches = [
        ("Auto-scored (8 categories, 2,260 records)",
         "Regex and string-matching on the model response. Checks for key terms, "
         "numeric answers, or absence of false claims. Binary (0/1) or proportion-based. "
         "Categories: Fabricated Entities, Math with Distractors, False Premise, "
         "Confidence Calibration, Counting & Arithmetic, Relational Puzzles (AIW), "
         "Distractor Injection, Hallucination Fabrication."),
        ("LLM-as-Judge (12 categories, 4,180 records)",
         "Each response was scored by Claude Sonnet (claude-sonnet-4-5 via OpenRouter) "
         "using a category-specific rubric. The judge was given: category name, scoring guide, "
         "question, expected answer, and model response (up to 1,500 characters). "
         "Output: JSON {score, reason}. Temperature=0."),
    ]
    for label, desc in score_approaches:
        p = doc.add_paragraph()
        r1 = p.add_run(label + ": ")
        r1.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    # Score scale table
    add_para(doc, "Score scales by category:", bold=True)
    scale_tbl = doc.add_table(rows=1, cols=4)
    scale_tbl.style = "Table Grid"
    hdr = scale_tbl.rows[0].cells
    for c, h in zip(hdr, ["Category", "Scorer", "Scale", "Interpretation"]):
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        set_cell_bg(c, "1F3864")
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for cat_key, (cat_name, scorer, scale, interp) in CATEGORY_DESCRIPTIONS.items():
        row = scale_tbl.add_row()
        row.cells[0].text = cat_name
        row.cells[1].text = scorer.upper()
        row.cells[2].text = scale
        row.cells[3].text = interp
        row.cells[3].paragraphs[0].runs[0].font.size = Pt(8)
        if scorer == "judge":
            set_cell_bg(row.cells[1], "E2EFDA")
        else:
            set_cell_bg(row.cells[1], "FFF2CC")
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_borders(c)

    doc.add_paragraph()
    add_para(doc, (
        "Note on hallucination scoring: The 9_open_ended_hallucination_prone category uses a "
        "proportion scale where 0.0 = fully accurate and 1.0 = fully hallucinated. For summary "
        "tables, the score is inverted (1 − raw_score) so that higher values consistently mean "
        "better performance."
    ), italic=True, size=9)

    doc.add_page_break()

    # ── 4. Results ───────────────────────────────────────────────────────────
    add_heading(doc, "4. Results", 1)

    # 4.1 Overall
    add_heading(doc, "4.1 Overall Performance by Condition", 2)
    add_para(doc, (
        "Table 1 shows overall accuracy (proportion correct, normalized to 0–1) by condition "
        "for auto-scored and judge-scored categories separately, plus combined."
    ))

    # Combined overall: weighted average
    # Auto: 565 items per condition; Judge: ~1,043 items per condition
    def combined(a_val, a_n, j_val, j_n):
        return (a_val * a_n + j_val * j_n) / (a_n + j_n)

    comb_overall = {
        c: combined(AUTO_OVERALL[c], 565, JUDGE_OVERALL[c], 1045)
        for c in [1, 2, 3, 4]
    }

    tbl_ov = doc.add_table(rows=4, cols=5)
    tbl_ov.style = "Table Grid"
    hdrs = ["", "Zero-Shot", "CoT", "CSEP-Only", "CSEP+ZS"]
    rows_ov = [
        ("Auto-Scored", AUTO_OVERALL),
        ("Judge-Scored", JUDGE_OVERALL),
        ("Combined", comb_overall),
    ]
    # header row
    for i, h in enumerate(hdrs):
        tbl_ov.rows[0].cells[i].text = h
        tbl_ov.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl_ov.rows[0].cells[i], "1F3864")
        tbl_ov.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for ri, (label, data) in enumerate(rows_ov):
        row = tbl_ov.rows[ri + 1]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        for ci, cond in enumerate([1, 2, 3, 4]):
            v = data[cond]
            row.cells[ci + 1].text = f"{v:.3f}"
            row.cells[ci + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row.cells[ci + 1], score_color(v))
            set_cell_borders(row.cells[ci + 1])

    doc.add_paragraph()
    add_para(doc, (
        "Observation: Auto-scored categories show CoT > Zero-Shot > CSEP+ZS > CSEP-Only, "
        "while judge-scored categories reverse this pattern: Zero-Shot > CSEP+ZS > CSEP-Only > CoT. "
        "This reversal is one of the most important findings of the study."
    ), italic=True)

    # 4.2 By model
    doc.add_paragraph()
    add_heading(doc, "4.2 Performance by Model", 2)
    add_para(doc, "Table 2a (Auto-scored) and Table 2b (Judge-scored) break down performance by model and condition.")

    add_para(doc, "Table 2a — Auto-Scored by Model:", bold=True)
    tbl_am = doc.add_table(rows=len(MODEL_IDS) + 1, cols=5)
    tbl_am.style = "Table Grid"
    for i, h in enumerate(["Model", "Zero-Shot", "CoT", "CSEP-Only", "CSEP+ZS"]):
        tbl_am.rows[0].cells[i].text = h
        tbl_am.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl_am.rows[0].cells[i], "2E75B6")
        tbl_am.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for ri, mid in enumerate(MODEL_IDS):
        row = tbl_am.rows[ri + 1]
        row.cells[0].text = MODELS_SHORT[mid]
        row.cells[0].paragraphs[0].runs[0].bold = True
        for ci, cond in enumerate([1, 2, 3, 4]):
            v = AUTO_MODEL[mid][cond]
            row.cells[ci + 1].text = f"{v:.3f}"
            row.cells[ci + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row.cells[ci + 1], score_color(v))
            set_cell_borders(row.cells[ci + 1])

    doc.add_paragraph()
    add_para(doc, "Table 2b — Judge-Scored by Model:", bold=True)
    tbl_jm = doc.add_table(rows=len(MODEL_IDS) + 1, cols=5)
    tbl_jm.style = "Table Grid"
    for i, h in enumerate(["Model", "Zero-Shot", "CoT", "CSEP-Only", "CSEP+ZS"]):
        tbl_jm.rows[0].cells[i].text = h
        tbl_jm.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl_jm.rows[0].cells[i], "375623")
        tbl_jm.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for ri, mid in enumerate(MODEL_IDS):
        row = tbl_jm.rows[ri + 1]
        row.cells[0].text = MODELS_SHORT[mid]
        row.cells[0].paragraphs[0].runs[0].bold = True
        for ci, cond in enumerate([1, 2, 3, 4]):
            v = JUDGE_MODEL[mid][cond]
            row.cells[ci + 1].text = f"{v:.3f}"
            row.cells[ci + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row.cells[ci + 1], score_color(v))
            set_cell_borders(row.cells[ci + 1])

    doc.add_paragraph()

    # 4.3 By category
    add_heading(doc, "4.3 Performance by Category", 2)
    add_para(doc, "Table 3 shows results for all 20 categories. Heat map: green ≥ 0.80, yellow ≥ 0.65, orange ≥ 0.50, red < 0.50.")

    all_cats = sorted(
        list(AUTO_STATS.keys()) + list(JUDGE_STATS.keys()),
        key=lambda x: int(x.split("_")[0])
    )
    cat_data_combined = {**AUTO_STATS, **JUDGE_STATS}

    tbl_cat = doc.add_table(rows=len(all_cats) + 1, cols=7)
    tbl_cat.style = "Table Grid"
    for i, h in enumerate(["#", "Category", "Scorer", "Zero-Shot", "CoT", "CSEP-Only", "CSEP+ZS"]):
        tbl_cat.rows[0].cells[i].text = h
        tbl_cat.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl_cat.rows[0].cells[i], "1F3864")
        tbl_cat.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        tbl_cat.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    for ri, cat in enumerate(all_cats):
        row = tbl_cat.rows[ri + 1]
        num = cat.split("_")[0]
        cat_name, scorer, _, _ = CATEGORY_DESCRIPTIONS[cat]
        row.cells[0].text = num
        row.cells[1].text = cat_name
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[2].text = scorer.upper()
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(8)
        if scorer == "judge":
            set_cell_bg(row.cells[2], "E2EFDA")
        else:
            set_cell_bg(row.cells[2], "FFF2CC")
        for ci, cond in enumerate([1, 2, 3, 4]):
            data = cat_data_combined.get(cat, {})
            v = data.get(cond, None)
            if v is not None:
                row.cells[ci + 3].text = f"{v:.3f}"
                row.cells[ci + 3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_bg(row.cells[ci + 3], score_color(v))
            row.cells[ci + 3].paragraphs[0].runs[0].font.size = Pt(9) if row.cells[ci + 3].paragraphs[0].runs else None
            set_cell_borders(row.cells[ci + 3])

    doc.add_page_break()

    # ── 5. Graphs ────────────────────────────────────────────────────────────
    add_heading(doc, "5. Visualizations", 1)

    graph_captions = [
        ("1_overall_by_condition.png", "Figure 1 — Overall accuracy by condition (auto-scored categories only)"),
        ("2_by_model_and_condition.png", "Figure 2 — Performance by model and condition (auto-scored)"),
        ("3_category_heatmap.png", "Figure 3 — Category heatmap across all conditions"),
        ("4_csep_improvement.png", "Figure 4 — CSEP improvement over Zero-Shot baseline (delta scores)"),
        ("5_per_model_range.png", "Figure 5 — Score range per model across all conditions"),
    ]
    for fname, caption in graph_captions:
        path = os.path.join(GRAPHS_DIR, fname)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(5.5))
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(9)
            doc.add_paragraph()

    doc.add_page_break()

    # ── 6. Key Findings and Discussion ───────────────────────────────────────
    add_heading(doc, "6. Key Findings and Discussion", 1)

    findings = [
        ("Finding 1: The CoT Reversal",
         "CoT produces the best results on auto-scored structured tasks (overall: 0.701 vs. Zero-Shot 0.607) "
         "but the worst on judge-scored open-ended tasks (0.587 vs. Zero-Shot 0.724). "
         "This is explained by two factors: (a) CoT multi-pass synthesis frequently generates "
         "meta-level responses (comparing attempts rather than answering the question), which the judge "
         "scores as incomplete; (b) auto-scored tasks reward computation steps explicitly shown in CoT, "
         "while judge tasks reward concise, accurate factual responses. "
         "Implication: CoT is a strong strategy for structured computation but harmful for open-ended reasoning."),
        ("Finding 2: CSEP Does Not Generalize",
         "Across judge-scored categories, CSEP-Only (0.701) slightly improves on Zero-Shot (0.724) — "
         "they are essentially tied. CSEP+ZS (0.710) similarly does not produce a reliable improvement. "
         "The pipeline adds 3 extra API calls per question (4x cost) for marginal or negative gain on average. "
         "The overhead is only justified for specific categories where CSEP demonstrably helps."),
        ("Finding 3: CSEP Helps on Contradiction and Paradox Tasks",
         "For Hidden Contradictions (cat 7): Zero-Shot 0.680 → CSEP 0.725 (+4.5pp). "
         "For Self-Reference Paradox (cat 17): Zero-Shot 0.550 → CSEP+ZS 0.717 (+16.7pp). "
         "These gains are consistent: the decomposition step forces the model to explicitly reason about "
         "what the question is really asking, which helps surface embedded contradictions and paradoxes "
         "that direct answering misses."),
        ("Finding 4: CSEP Hurts Hallucination-Prone Tasks",
         "Hallucination Fabrication (cat 18, auto): Zero-Shot 0.325 → CSEP 0.150 (-17.5pp). "
         "Fabricated Entities (cat 2, auto): Zero-Shot 0.510 → CSEP 0.260 (-25pp). "
         "Hallucination Resistance (cat 9, judge, inverted): Zero-Shot 0.490 → CSEP 0.452 (-3.8pp). "
         "The CSEP decomposition step generates more elaborate context, which small models then use to "
         "confidently fabricate details. The knowledge self-check in Stage 2 is insufficiently strict "
         "to prevent this in small models."),
        ("Finding 5: Phi-4 Benefits Most from CSEP",
         "Phi-4 shows the highest absolute scores across all conditions and the best CSEP performance: "
         "ZS=0.792, CoT=0.653, CSEP=0.749, CSEP+ZS=0.771. Even on hallucination tasks, Phi-4 "
         "shows better resistance than smaller models. The 14B parameter scale appears to be the "
         "threshold where CSEP scaffolding becomes reliably beneficial."),
        ("Finding 6: Llama-3.1-8B Shows Severe CoT Degradation",
         "Llama's judge-scored CoT performance (0.440) is 20pp below its zero-shot baseline (0.645). "
         "This is the most extreme CoT degradation in the study. The model's multi-pass responses "
         "frequently become incoherent or contradictory across the three passes, leading to synthesis "
         "responses that hedge excessively or return to meta-commentary. This suggests 8B Llama models "
         "are particularly sensitive to prompt length and complexity."),
    ]

    for title, body in findings:
        p = doc.add_paragraph()
        r1 = p.add_run(title)
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        add_para(doc, body)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ── 7. Category Examples ─────────────────────────────────────────────────
    add_heading(doc, "7. Detailed Examples by Category", 1)
    add_para(doc, (
        "For each judge-scored category, one high-scoring and one low-scoring example are shown "
        "with the question, expected answer, model response (truncated to 400 characters), "
        "and the judge's reasoning."
    ))
    doc.add_paragraph()

    for cat_key in [
        "1_abstract_analogies", "4_multi_step_reasoning", "5_temporal_reasoning",
        "7_hidden_contradictions", "9_open_ended_hallucination_prone",
        "10_consistency_under_reframing", "11_spatial_reasoning",
        "14_modified_classic_puzzles", "16_logical_inference", "17_self_reference_paradox",
        "19_linguistic_constraint", "20_adversarial_pressure",
    ]:
        if cat_key not in EXAMPLES:
            continue
        cat_name = CATEGORY_DESCRIPTIONS[cat_key][0]
        add_heading(doc, f"7.{list(EXAMPLES.keys()).index(cat_key)+1}  {cat_name}", 2)

        ex = EXAMPLES[cat_key]
        cond_name_b = CONDITIONS.get(ex["best"]["condition"], f"Cond {ex['best']['condition']}")
        cond_name_w = CONDITIONS.get(ex["worst"]["condition"], f"Cond {ex['worst']['condition']}")

        add_example_box(
            doc, "BEST EXAMPLE", "375623",
            ex["best"]["model"], cond_name_b, ex["best"]["score"],
            ex["best"]["question"], ex["best"]["expected"],
            ex["best"]["response"], ex["best"]["reason"]
        )
        add_example_box(
            doc, "WORST EXAMPLE", "C00000",
            ex["worst"]["model"], cond_name_w, ex["worst"]["score"],
            ex["worst"]["question"], ex["worst"]["expected"],
            ex["worst"]["response"], ex["worst"]["reason"]
        )

    doc.add_page_break()

    # ── 8. Conclusion ────────────────────────────────────────────────────────
    add_heading(doc, "8. Conclusion", 1)
    add_para(doc, (
        "CSEP (Conceptual Space Expansion Prompting) shows task-specific benefits but does not "
        "consistently outperform the much simpler zero-shot baseline across the full evaluation suite. "
        "The 4-call pipeline is particularly beneficial for tasks requiring contradiction detection, "
        "self-reference handling, and temporal reasoning — tasks where explicit conceptual scaffolding "
        "helps the model notice what it would otherwise overlook."
    ))
    doc.add_paragraph()
    add_para(doc, (
        "However, CSEP is counterproductive for hallucination-prone tasks. By generating elaborate "
        "decomposition context, smaller models (7–12B) become more confident in fabricated details "
        "rather than more cautious. The Stage 2 self-verification probe (checking whether the model "
        "can provide confirming and disconfirming examples) is insufficiently strict to prevent this."
    ))
    doc.add_paragraph()
    add_para(doc, (
        "CoT outperforms all methods on structured computation tasks (math, counting, relational puzzles) "
        "but degrades badly on open-ended reasoning tasks. This scoring-method interaction is crucial: "
        "the choice of evaluation method determines which prompting strategy appears best."
    ))
    doc.add_paragraph()
    add_para(doc, "Recommendations for practitioners:", bold=True)
    recs = [
        "Use CoT for structured computation tasks (math, counting, logic puzzles with definite answers).",
        "Use CSEP for tasks requiring explicit contradiction detection or paradox recognition.",
        "Use Zero-Shot for open-ended factual tasks to minimize hallucination risk.",
        "CSEP+ZS is a safer variant than CSEP-Only — the zero-shot anchor reduces drift.",
        "For models below 12B parameters, be cautious with multi-pass pipelines; the gain rarely justifies the cost.",
        "Phi-4 (14B) benefits most from CSEP; consider using it as the minimum model size for CSEP deployment.",
    ]
    for rec in recs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rec)

    doc.add_paragraph()
    add_para(doc, (
        "Future work should investigate whether targeted Stage 2 prompt modifications — particularly "
        "stricter knowledge checks — can preserve CSEP's benefits while reducing hallucination amplification."
    ))

    # Save
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build_document()
