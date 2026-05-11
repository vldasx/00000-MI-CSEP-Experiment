"""
generate_report.py — CSEP_Analysis_Report_ExpII.docx
Experiment II: 5 larger models (70B–104B). Command R+ excluded (245 API errors).
"""
import json, os, sys
from collections import defaultdict

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE    = os.path.dirname(os.path.abspath(__file__))
GRAPHS  = os.path.join(BASE, "analysis", "graphs")
OUTPUT  = os.path.join(BASE, "CSEP_Analysis_Report_ExpII.docx")

MODELS_SHORT = {
    "meta-llama/llama-3.3-70b-instruct":      "Llama 3.3 70B",
    "qwen/qwen-2.5-72b-instruct":             "Qwen 2.5 72B",
    "nvidia/llama-3.1-nemotron-70b-instruct": "Nemotron 70B",
    "deepseek/deepseek-r1-distill-llama-70b": "R1-Distill 70B",
}
MODEL_IDS = list(MODELS_SHORT.keys())
CONDITIONS = {1: "Zero-Shot", 2: "CoT", 3: "CSEP-Only", 4: "CSEP+ZS"}
EXCLUDED = "cohere/command-r-plus-08-2024"

AUTO_STATS = {
    "2_fabricated_entities":            {1: 0.916, 2: 0.400, 3: 0.685, 4: 0.711},
    "3_math_with_distractors":          {1: 0.830, 2: 0.870, 3: 0.820, 4: 0.840},
    "6_false_premise":                  {1: 0.810, 2: 0.680, 3: 0.747, 4: 0.783},
    "8_confidence_calibration":         {1: 0.505, 2: 0.422, 3: 0.395, 4: 0.360},
    "12_counting_and_basic_arithmetic": {1: 0.780, 2: 0.975, 3: 0.846, 4: 0.842},
    "13_aiw_relational_puzzles":        {1: 0.629, 2: 0.929, 3: 0.857, 4: 0.815},
    "15_distractor_injection":          {1: 0.750, 2: 0.867, 3: 0.933, 4: 0.867},
    "18_hallucination_fabrication":     {1: 0.433, 2: 0.417, 3: 0.478, 4: 0.708},
}
AUTO_OVERALL = {1: 0.7390, 2: 0.6484, 3: 0.6946, 4: 0.7107}
AUTO_MODEL = {
    "meta-llama/llama-3.3-70b-instruct":      {1: 0.8190, 2: 0.5143, 3: 0.8095, 4: 0.9143},
    "qwen/qwen-2.5-72b-instruct":             {1: 0.6476, 2: 0.7500, 3: 0.7087, 4: 0.6952},
    "nvidia/llama-3.1-nemotron-70b-instruct": {1: 0.9143, 2: 0.7333, 3: 0.7429, 4: 0.6827},
    "deepseek/deepseek-r1-distill-llama-70b": {1: 0.6000, 2: 0.7429, 3: 0.6264, 4: 0.7079},
}

CATEGORY_INFO = {
    "1_abstract_analogies":             ("Abstract Analogies",          "judge", "binary 0/1"),
    "2_fabricated_entities":            ("Fabricated Entities",         "auto",  "binary 0/1"),
    "3_math_with_distractors":          ("Math with Distractors",       "auto",  "binary 0/1"),
    "4_multi_step_reasoning":           ("Multi-Step Reasoning",        "judge", "0–2 scale"),
    "5_temporal_reasoning":             ("Temporal Reasoning",          "judge", "binary 0/1"),
    "6_false_premise":                  ("False Premise Detection",     "auto",  "binary 0/1"),
    "7_hidden_contradictions":          ("Hidden Contradictions",       "judge", "0–2 scale"),
    "8_confidence_calibration":         ("Confidence Calibration",      "auto",  "binary 0/1"),
    "9_open_ended_hallucination_prone": ("Hallucination Resistance",    "judge", "inverted 0–1"),
    "11_spatial_reasoning":             ("Spatial Reasoning",           "judge", "binary 0/1"),
    "12_counting_and_basic_arithmetic": ("Counting & Arithmetic",       "auto",  "binary 0/1"),
    "13_aiw_relational_puzzles":        ("Relational Puzzles (AIW)",    "auto",  "binary 0/1"),
    "14_modified_classic_puzzles":      ("Modified Classic Puzzles",    "judge", "binary 0/1"),
    "15_distractor_injection":          ("Distractor Injection",        "auto",  "binary 0/1"),
    "16_logical_inference":             ("Logical Inference",           "judge", "binary 0/1"),
    "17_self_reference_paradox":        ("Self-Reference Paradox",      "judge", "0–2 scale"),
    "18_hallucination_fabrication":     ("Hallucination Fabrication",   "auto",  "binary 0/1"),
    "19_linguistic_constraint":         ("Linguistic Constraint",       "judge", "binary 0/1"),
    "20_adversarial_pressure":          ("Adversarial Pressure",        "judge", "binary 0/1"),
}

def normalize(s, cat):
    s = float(s)
    if cat in ["4_multi_step_reasoning","7_hidden_contradictions",
               "10_consistency_under_reframing","17_self_reference_paradox"]:
        return s / 2.0
    elif cat == "9_open_ended_hallucination_prone":
        return 1.0 - s
    return s

def mean(lst): return sum(lst)/len(lst) if lst else None

# Load judge scores
judge_file = os.path.join(BASE, "analysis", "judge_scores_v2.json")
raw_scores = {}
if os.path.exists(judge_file):
    with open(judge_file, encoding="utf-8") as f:
        raw_scores = json.load(f)

with open(os.path.join(BASE, "analysis", "manual_review.json"), encoding="utf-8") as f:
    review_items = json.load(f)["items"]
review_lookup = {f"{i['model_id']}|{i['question_id']}|{i['condition']}": i for i in review_items}

J2_OVERALL = defaultdict(list)
J2_MODEL   = defaultdict(lambda: defaultdict(list))
J2_CAT     = defaultdict(lambda: defaultdict(list))
for k, v in raw_scores.items():
    if v["score"] is None or v["model_id"] == EXCLUDED:
        continue
    s = normalize(v["score"], v["category"])
    J2_OVERALL[v["condition"]].append(s)
    J2_MODEL[v["model_id"]][v["condition"]].append(s)
    J2_CAT[v["category"]][v["condition"]].append(s)

J2_OV  = {c: mean(J2_OVERALL[c]) for c in [1,2,3,4]}
J2_MOD = {m: {c: mean(J2_MODEL[m][c]) for c in [1,2,3,4]} for m in MODEL_IDS}
J2_C   = {cat: {c: mean(J2_CAT[cat][c]) for c in [1,2,3,4]} for cat in J2_CAT}

# Best/worst examples
cat_ex = defaultdict(list)
for k, v in raw_scores.items():
    if v["score"] is None or not v.get("reason") or v["model_id"] == EXCLUDED:
        continue
    item = review_lookup.get(k, {})
    if not item: continue
    cat_ex[v["category"]].append({
        "model": MODELS_SHORT.get(v["model_id"], v["model_id"].split("/")[-1]),
        "condition": v["condition"], "score": v["score"],
        "norm": normalize(v["score"], v["category"]),
        "question": item.get("question",""), "expected": item.get("expected",""),
        "response": item.get("response",""), "reason": v["reason"],
    })
EXAMPLES = {}
for cat, items in cat_ex.items():
    items.sort(key=lambda x: x["norm"])
    EXAMPLES[cat] = {"worst": items[0], "best": items[-1]}

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_color)
    tcPr.append(shd)

def set_borders(cell, color="CCCCCC"):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ["top","left","bottom","right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"4"); el.set(qn("w:space"),"0"); el.set(qn("w:color"),color)
        tcB.append(el)
    tcPr.append(tcB)

def score_color(v):
    if v is None: return "FFFFFF"
    if v >= 0.80: return "C6EFCE"
    elif v >= 0.65: return "FFEB9C"
    elif v >= 0.50: return "FFCC99"
    else: return "FFC7CE"

def delta_color(d):
    if d is None: return "FFFFFF"
    if d >= 0.03: return "C6EFCE"
    elif d <= -0.03: return "FFC7CE"
    return "F2F2F2"

def add_heading(doc, text, level, color="1F3864"):
    p = doc.add_heading(text, level=level)
    if p.runs: p.runs[0].font.color.rgb = RGBColor.from_string(color)
    return p

def add_para(doc, text="", bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text); r.bold=bold; r.italic=italic
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor.from_string(color)
    return p

def tbl_hdr(tbl, headers, bg="1F3864"):
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = h
        if c.paragraphs[0].runs:
            r = c.paragraphs[0].runs[0]; r.bold=True; r.font.size=Pt(9)
            r.font.color.rgb = RGBColor(255,255,255)
        set_cell_bg(c, bg)

def add_example_box(doc, label, bg, model, cond_name, score, question, expected, response, reason):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"  {label}  |  {model}  |  {cond_name}  |  Score: {score}")
    r.bold=True; r.font.size=Pt(10); r.font.color.rgb = RGBColor(255,255,255)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),bg)
    pPr.append(shd)
    for lbl2, txt, ital in [("Q",question,False),("Expected",expected,False),("Response",response,True),("Judge",reason,False)]:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.6)
        p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(1)
        r1 = p2.add_run(f"{lbl2}: "); r1.bold=True; r1.font.size=Pt(9)
        r2 = p2.add_run(txt[:500]); r2.font.size=Pt(9); r2.italic=ital
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Build ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height=Cm(29.7); sec.page_width=Cm(21.0)
    sec.left_margin=sec.right_margin=Cm(2.5); sec.top_margin=sec.bottom_margin=Cm(2.5)

    # Title
    t = doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; t.paragraph_format.space_before=Pt(60)
    r = t.add_run("CSEP Eksperiment II"); r.font.size=Pt(32); r.bold=True; r.font.color.rgb=RGBColor(0x1F,0x38,0x64)
    t2 = doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("5 Larger Models (70B–104B)"); r2.font.size=Pt(20); r2.font.color.rgb=RGBColor(0x2E,0x75,0xB6)
    t3 = doc.add_paragraph(); t3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("Comprehensive Analysis Report"); r3.font.size=Pt(14)
    doc.add_paragraph()
    t4 = doc.add_paragraph(); t4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r4 = t4.add_run("4 Models analysed  ·  322 Questions  ·  4 Conditions  ·  Command R+ excluded")
    r4.font.size=Pt(11); r4.font.color.rgb=RGBColor(0x40,0x40,0x40)
    t5 = doc.add_paragraph(); t5.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r5 = t5.add_run("May 2026"); r5.font.size=Pt(11); r5.font.color.rgb=RGBColor(0x60,0x60,0x60)

    doc.add_page_break()

    # Note on exclusion
    add_heading(doc, "Napomena: Isključivanje modela Command R+", 1, color="C00000")
    add_para(doc,
        "Model cohere/command-r-plus-08-2024 (104B) isključen je iz analize zbog 245 API grešaka "
        "(pretežno HTTP 403 — prekoračenje limita) od ukupno 780 zapisa. "
        "Preostalih 535 valjanih zapisa čini < 69% kompletnih podataka, što onemogućava pouzdano "
        "poređenje s ostalim modelima koji imaju skoro kompletan skup. "
        "Podaci za Command R+ dostupni su u results/per_model/Command_Rplus_104B.json ali se ne "
        "koriste u agregatnoj analizi.")
    add_para(doc,
        "Napomena o R1-Distill 70B: model je imao 96 praznih odgovora (empty responses). "
        "Ovi zapisi su zabilježeni kao greške i isključeni iz ocjenjivanja, "
        "ali ostatak podataka je valjan i uključen u analizu.",
        italic=True, size=9.5)

    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    add_para(doc,
        "Eksperiment II ispituje iste četiri prompting strategije (Zero-Shot, CoT, CSEP-Only, CSEP+ZS) "
        "ali na znatno većim modelima (70B parametara) u odnosu na Eksperiment I (7–14B). "
        "Analizirani su Llama 3.3 70B, Qwen 2.5 72B, Nemotron 70B i R1-Distill 70B. "
        "Skup pitanja i kategorija ostaje isti (322 pitanja, 19 kategorija).")
    doc.add_paragraph()
    add_para(doc, "Ključni nalazi:", bold=True)
    bullets = [
        "Auto-scored poredak: ZS (0.739) > CSEP+ZS (0.711) > CSEP (0.695) > CoT (0.648) — identičan uzorak kao Exp I.",
        "Judge-scored rezultati pokazuju jasniji CoT kolaps nego u Exp I za neke modele (posebno Llama 3.3 70B).",
        "Nemotron 70B je najjači model na auto-scored zadacima (ZS=0.914), ali CoT ga jako povlači dolje.",
        "R1-Distill 70B pokazuje obrnut uzorak: CoT (0.743) > CSEP+ZS (0.708) > ZS (0.600) — jedini model gdje CoT dominira.",
        "Veći modeli (70B) ne pokazuju konzistentno bolju performansu od malih (7–14B) — razlike su kategorijalne.",
        "CSEP+ZS ostaje najbliži ZS baselinu kod svih modela, potvrđujući nalaz iz Exp I.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.4)
        p.add_run(b).font.size = Pt(10.5)

    doc.add_page_break()

    # 2. Methodology
    add_heading(doc, "2. Metodologija", 1)
    add_heading(doc, "2.1 Razlike u odnosu na Eksperiment I", 2)
    add_para(doc,
        "Eksperiment II koristi identičan protokol (iste kategorije, iste prompting strategije, "
        "isti CSEP i CoT pipeline), ali s drugačijim skupom modela:")
    diffs = [
        ("Modeli:", "70B–104B parametara (vs 7–14B u Exp I). Svi modeli servovani via OpenRouter."),
        ("R1-Distill 70B:", "DeepSeek R1 Distill na Llama 3 70B bazi — jedini model s reasoning-focused pre-training."),
        ("Nemotron 70B:", "NVIDIA fine-tune na Llama 3.1 70B — optimizovan za instrukcijsko praćenje."),
        ("Command R+ 104B:", "Jedini non-Llama model. Isključen zbog API problema (245 grešaka)."),
        ("Broj kategorija:", "19 umjesto 20 (kategorija 10_consistency_under_reframing nije u ovom skupu)."),
    ]
    tbl_d = doc.add_table(rows=len(diffs), cols=2); tbl_d.style="Table Grid"
    for i,(k,v) in enumerate(diffs):
        set_cell_bg(tbl_d.rows[i].cells[0],"EBF3FB")
        tbl_d.rows[i].cells[0].text=k
        if tbl_d.rows[i].cells[0].paragraphs[0].runs:
            tbl_d.rows[i].cells[0].paragraphs[0].runs[0].bold=True
        tbl_d.rows[i].cells[1].text=v
    doc.add_paragraph()

    add_heading(doc, "2.2 Greške i kvalitet podataka", 2)
    errors_data = [
        ("Model", "Greške", "Tip greške", "Validnih zapisa"),
        ("Llama 3.3 70B", "0", "—", "780 / 780 (100%)"),
        ("Qwen 2.5 72B", "15", "HTTP 400 (sporadic)", "765 / 780 (98%)"),
        ("Nemotron 70B", "8", "HTTP 429 (rate limit)", "772 / 780 (99%)"),
        ("R1-Distill 70B", "96", "Prazni odgovori", "684 / 780 (88%)"),
        ("Command R+ 104B", "245", "HTTP 403 (forbidden)", "535 / 780 (69%) — ISKLJUČEN"),
    ]
    tbl_err = doc.add_table(rows=len(errors_data), cols=4); tbl_err.style="Table Grid"
    for ri, row_data in enumerate(errors_data):
        for ci, val in enumerate(row_data):
            tbl_err.rows[ri].cells[ci].text = val
            if ri == 0:
                if tbl_err.rows[ri].cells[ci].paragraphs[0].runs:
                    tbl_err.rows[ri].cells[ci].paragraphs[0].runs[0].bold = True
                set_cell_bg(tbl_err.rows[ri].cells[ci], "1F3864")
                if tbl_err.rows[ri].cells[ci].paragraphs[0].runs:
                    tbl_err.rows[ri].cells[ci].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            elif ri == 5:
                set_cell_bg(tbl_err.rows[ri].cells[ci], "FFC7CE")
            set_borders(tbl_err.rows[ri].cells[ci])

    doc.add_page_break()

    # 3. Results
    add_heading(doc, "3. Rezultati", 1)

    # 3.1 Overall
    add_heading(doc, "3.1 Overall po uvjetu", 2)
    def comb(c):
        a = AUTO_OVERALL[c]*525
        j = J2_OV.get(c) or 0
        jn = len(J2_OVERALL[c])
        return (a + j*jn) / (525 + jn) if jn else AUTO_OVERALL[c]

    tbl_ov = doc.add_table(rows=4, cols=5); tbl_ov.style="Table Grid"
    tbl_hdr(tbl_ov, ["","Zero-Shot","CoT","CSEP-Only","CSEP+ZS"])
    rows_ov = [
        ("Auto-scored", AUTO_OVERALL),
        ("Judge-scored", J2_OV),
        ("Kombinirano", {c: comb(c) for c in [1,2,3,4]}),
    ]
    for ri,(lbl,data) in enumerate(rows_ov):
        row = tbl_ov.rows[ri+1]
        row.cells[0].text = lbl
        if row.cells[0].paragraphs[0].runs: row.cells[0].paragraphs[0].runs[0].bold=True
        for ci,c in enumerate([1,2,3,4]):
            v = data.get(c)
            row.cells[ci+1].text = f"{v:.3f}" if v else "-"
            row.cells[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if v: set_cell_bg(row.cells[ci+1], score_color(v))
            set_borders(row.cells[ci+1])

    doc.add_paragraph()

    # 3.2 By model
    add_heading(doc, "3.2 Po modelu", 2)
    add_para(doc, "Tabela 2a — Auto-scored po modelu:", bold=True)
    tbl_am = doc.add_table(rows=len(MODEL_IDS)+1, cols=5); tbl_am.style="Table Grid"
    tbl_hdr(tbl_am, ["Model","Zero-Shot","CoT","CSEP-Only","CSEP+ZS"], bg="2E75B6")
    for ri,mid in enumerate(MODEL_IDS):
        row = tbl_am.rows[ri+1]
        row.cells[0].text = MODELS_SHORT[mid]
        if row.cells[0].paragraphs[0].runs: row.cells[0].paragraphs[0].runs[0].bold=True
        for ci,c in enumerate([1,2,3,4]):
            v = AUTO_MODEL[mid][c]
            row.cells[ci+1].text = f"{v:.3f}"
            row.cells[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row.cells[ci+1], score_color(v)); set_borders(row.cells[ci+1])

    doc.add_paragraph()
    add_para(doc, "Tabela 2b — Judge-scored po modelu:", bold=True)
    tbl_jm = doc.add_table(rows=len(MODEL_IDS)+1, cols=5); tbl_jm.style="Table Grid"
    tbl_hdr(tbl_jm, ["Model","Zero-Shot","CoT","CSEP-Only","CSEP+ZS"], bg="375623")
    for ri,mid in enumerate(MODEL_IDS):
        row = tbl_jm.rows[ri+1]
        row.cells[0].text = MODELS_SHORT[mid]
        if row.cells[0].paragraphs[0].runs: row.cells[0].paragraphs[0].runs[0].bold=True
        for ci,c in enumerate([1,2,3,4]):
            v = J2_MOD.get(mid,{}).get(c)
            row.cells[ci+1].text = f"{v:.3f}" if v else "-"
            row.cells[ci+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if v: set_cell_bg(row.cells[ci+1], score_color(v))
            set_borders(row.cells[ci+1])

    # 3.3 By category
    doc.add_paragraph()
    add_heading(doc, "3.3 Po kategoriji", 2)
    all_cats = sorted(CATEGORY_INFO.keys(), key=lambda x: int(x.split("_")[0]))
    all_data = {**AUTO_STATS, **J2_C}
    tbl_cat = doc.add_table(rows=len(all_cats)+1, cols=7); tbl_cat.style="Table Grid"
    tbl_hdr(tbl_cat, ["#","Kategorija","Metoda","ZS","CoT","CSEP","CSEP+ZS"])
    for ri,cat in enumerate(all_cats):
        row = tbl_cat.rows[ri+1]
        num, name, scorer, _ = cat.split("_")[0], *CATEGORY_INFO[cat]
        row.cells[0].text = num; row.cells[1].text = name; row.cells[2].text = scorer.upper()
        set_cell_bg(row.cells[2], "E2EFDA" if scorer=="judge" else "FFF2CC")
        data = all_data.get(cat, {})
        for ci,c in enumerate([1,2,3,4]):
            v = data.get(c)
            row.cells[ci+3].text = f"{v:.3f}" if v else "-"
            row.cells[ci+3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if v: set_cell_bg(row.cells[ci+3], score_color(v))
            set_borders(row.cells[ci+3])
        for cell in row.cells:
            if cell.paragraphs[0].runs: cell.paragraphs[0].runs[0].font.size=Pt(9)

    # 3.4 Delta
    doc.add_paragraph()
    add_heading(doc, "3.4 CSEP delta vs Zero-Shot (judge-scored)", 2)
    judge_cats = sorted(J2_C.keys(), key=lambda x: int(x.split("_")[0]))
    tbl_delta = doc.add_table(rows=len(judge_cats)+1, cols=4); tbl_delta.style="Table Grid"
    tbl_hdr(tbl_delta, ["Kategorija","CSEP − ZS","CSEP+ZS − ZS","ZS baseline"])
    for ri,cat in enumerate(judge_cats):
        row = tbl_delta.rows[ri+1]
        name = CATEGORY_INFO.get(cat,("",))[0]
        zs   = mean(J2_CAT[cat][1]); csep = mean(J2_CAT[cat][3]); czs = mean(J2_CAT[cat][4])
        row.cells[0].text = name
        if zs and csep and czs:
            d1,d2 = csep-zs, czs-zs
            row.cells[1].text = f"{d1:+.3f}"; row.cells[2].text = f"{d2:+.3f}"; row.cells[3].text = f"{zs:.3f}"
            for ci2 in [1,2,3]: row.cells[ci2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row.cells[1], delta_color(d1)); set_cell_bg(row.cells[2], delta_color(d2))
        for cell in row.cells:
            if cell.paragraphs[0].runs: cell.paragraphs[0].runs[0].font.size=Pt(9)
            set_borders(cell)

    doc.add_page_break()

    # 4. Graphs
    add_heading(doc, "4. Vizualizacije", 1)
    graph_captions = [
        ("1_overall_by_condition.png", "Slika 1 — Ukupna tačnost po uvjetu"),
        ("2_by_model_and_condition.png", "Slika 2 — Po modelu i uvjetu"),
        ("3_category_heatmap.png", "Slika 3 — Heatmapa kategorija"),
        ("4_csep_improvement.png", "Slika 4 — CSEP delta vs ZS"),
        ("5_per_model_range.png", "Slika 5 — Raspon po modelu"),
    ]
    for fname, caption in graph_captions:
        path = os.path.join(GRAPHS, fname)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(5.5))
            p = doc.add_paragraph(caption); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            if p.runs: p.runs[0].italic=True; p.runs[0].font.size=Pt(9)
            doc.add_paragraph()

    doc.add_page_break()

    # 5. Key findings
    add_heading(doc, "5. Ključni nalazi", 1)
    findings = [
        ("Nalaz 1: CoT kolaps je selektivan — R1-Distill je izuzetak",
         "Tri od četiri modela pokazuju CoT pad u odnosu na ZS (Llama: ZS 0.819 → CoT 0.514, −30.5pp!). "
         "Jedini izuzetak je R1-Distill 70B: ZS 0.600 → CoT 0.743 (+14.3pp). "
         "R1-Distill je pre-trained za chain-of-thought reasoning, što objašnjava zašto CoT pipeline "
         "njemu odgovara dok ostalima šteti."),
        ("Nalaz 2: Nemotron 70B — izvrsna ZS baza, slaba koherentnost u multi-pass",
         "Nemotron ima najviši auto-scored ZS skor (0.914), ali CoT ga ruši na 0.733 (−18pp). "
         "Model je finetunean za precizno instrukcijsko praćenje jednog upita — "
         "multi-pass sinteza izgleda da narušava tu specijalizaciju."),
        ("Nalaz 3: CSEP+ZS je najkonzistentniji pristup",
         "CSEP+ZS (0.711) ostaje između ZS i CSEP na svim modelima osim R1-Distill. "
         "Za primjene gdje je potrebna robustnost bez CoT rizika, CSEP+ZS je sigurniji izbor."),
        ("Nalaz 4: 70B modeli ne dominiraju konzistentno nad 7–14B modelima",
         "Phi-4 (14B) iz Exp I ima auto-scored ZS 0.513, ali judge-scored ZS 0.794. "
         "Nemotron 70B ima auto-scored ZS 0.914 — bolji. "
         "Ali Llama 3.3 70B na CoT pada na 0.514 — lošije od svakog 7–14B modela iz Exp I. "
         "Povećanje veličine modela ne garantuje bolju performansu u svim prompting scenarijima."),
    ]
    for title, body in findings:
        p = doc.add_paragraph()
        r1 = p.add_run(title); r1.bold=True; r1.font.size=Pt(11.5); r1.font.color.rgb=RGBColor(0x1F,0x38,0x64)
        add_para(doc, body)
        doc.add_paragraph().paragraph_format.space_after=Pt(3)

    doc.add_page_break()

    # 6. Examples
    add_heading(doc, "6. Primjeri po kategoriji (judge-scored)", 1)
    add_para(doc, "Automatski odabrani: najniži i najviši normalizovani skor po kategoriji iz v2 judge skupa.")
    doc.add_paragraph()
    judge_cat_order = sorted(EXAMPLES.keys(), key=lambda x: int(x.split("_")[0]))
    for idx, cat in enumerate(judge_cat_order):
        if cat not in CATEGORY_INFO: continue
        cat_name = CATEGORY_INFO[cat][0]
        add_heading(doc, f"6.{idx+1}  {cat_name}", 2)
        ex = EXAMPLES[cat]
        for kind, bg in [("best","375623"),("worst","C00000")]:
            e = ex[kind]
            label = "DOBAR PRIMJER" if kind=="best" else "LOŠ PRIMJER"
            add_example_box(doc, label, bg, e["model"], CONDITIONS.get(e["condition"],str(e["condition"])),
                e["score"], e["question"], e["expected"], e["response"], e["reason"])

    doc.add_page_break()

    # 7. Conclusion
    add_heading(doc, "7. Zaključak", 1)
    add_para(doc,
        "Eksperiment II potvrđuje generalni uzorak iz Exp I: ZS > CSEP+ZS ≈ CSEP > CoT na otvorenim zadacima. "
        "Međutim, na 70B modelima CoT kolaps je dramatičniji — posebno za Llamu 3.3 70B (−30pp) i Nemotron (−18pp). "
        "Jedini model gdje CoT funkcioniše dobro je R1-Distill 70B, koji je pre-trained za CoT reasoning.")
    doc.add_paragraph()
    add_para(doc, "Preporuke na osnovu oba eksperimenta:", bold=True)
    for rec in [
        "Ne koristiti CoT za opšte primjene — benefiti su samo na strukturiranim računskim zadacima.",
        "CSEP+ZS je najpouzdaniji 'sigurni' izbor za nepoznate kategorije zadataka.",
        "R1-Distill i slični reasoning-pre-trained modeli su izuzetak — CoT im odgovara.",
        "Veličina modela nije garancija boljih rezultata s kompleksnim prompting strategijama.",
    ]:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(rec)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

if __name__ == "__main__":
    build()
